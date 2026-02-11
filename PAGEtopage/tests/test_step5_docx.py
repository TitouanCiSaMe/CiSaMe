"""
Tests unitaires pour l'etape 5: Conversion DOCX -> Vertical
"""

import pytest
from pathlib import Path
import tempfile

from docx import Document
from docx.shared import RGBColor, Pt

from PAGEtopage.config import Config
from PAGEtopage.models import ExtractedPage, PageMetadata
from PAGEtopage.step5_docx.docx_parser import (
    DocxParser, HEADER_MARKERS, PAGE_HEADER_PATTERN, _parse_page_header,
)


def _create_latin_analyzer_docx(path: Path, text_lines: list[str]) -> None:
    """
    Cree un fichier DOCX imitant la sortie de latin_analyzer

    Args:
        path: Chemin de sortie
        text_lines: Lignes de texte a inclure
    """
    doc = Document()

    # En-tete latin_analyzer
    title = doc.add_paragraph()
    run = title.add_run("Analyse de texte latin medieval")
    run.bold = True
    run.font.size = Pt(14)

    legend = doc.add_paragraph()
    legend.add_run("Légende : ").bold = True
    black_run = legend.add_run("Noir = OK (score ≥75) ")
    black_run.font.color.rgb = RGBColor(0, 0, 0)
    orange_run = legend.add_run("Orange = À vérifier (score 40-74) ")
    orange_run.font.color.rgb = RGBColor(255, 140, 0)
    red_run = legend.add_run("Rouge = Erreur probable (score <40)")
    red_run.font.color.rgb = RGBColor(255, 0, 0)

    # Separateur
    doc.add_paragraph("_" * 80)

    # Texte
    for line in text_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(10)

    doc.save(str(path))


def _create_multipage_docx(path: Path, pages: list[dict]) -> None:
    """
    Cree un DOCX multi-pages imitant latin_analyzer v2.4+

    Args:
        path: Chemin de sortie
        pages: Liste de dicts {folio, page_number, running_title, lines}
    """
    doc = Document()

    # En-tete latin_analyzer
    title = doc.add_paragraph()
    run = title.add_run("Analyse de texte latin medieval")
    run.bold = True
    run.font.size = Pt(14)

    legend = doc.add_paragraph()
    legend.add_run("Légende : ").bold = True
    black_run = legend.add_run("Noir = OK (score ≥75) ")
    black_run.font.color.rgb = RGBColor(0, 0, 0)
    orange_run = legend.add_run("Orange = À vérifier (score 40-74) ")
    orange_run.font.color.rgb = RGBColor(255, 140, 0)
    red_run = legend.add_run("Rouge = Erreur probable (score <40)")
    red_run.font.color.rgb = RGBColor(255, 0, 0)

    # Separateur
    doc.add_paragraph("_" * 80)

    # Pages avec en-tetes
    for page_info in pages:
        # En-tete de page
        parts = []
        if page_info.get("folio"):
            parts.append(f"Folio: {page_info['folio']}")
        if page_info.get("page_number") is not None:
            parts.append(f"Page: {page_info['page_number']}")
        if page_info.get("running_title"):
            parts.append(page_info["running_title"])
        header_text = " | ".join(parts) if parts else "Page"
        header_line = f"{'─' * 20} {header_text} {'─' * 20}"

        p = doc.add_paragraph()
        run = p.add_run(header_line)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.bold = True

        # Lignes de texte
        for line in page_info.get("lines", []):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(10)

    doc.save(str(path))


def _create_plain_docx(path: Path, text_lines: list[str]) -> None:
    """Cree un DOCX simple sans en-tete latin_analyzer"""
    doc = Document()
    for line in text_lines:
        doc.add_paragraph(line)
    doc.save(str(path))


def _make_config(**corpus_kwargs) -> Config:
    """Cree une config avec metadonnees de corpus"""
    config = Config()
    config.enrichment.lemmatizer = "simple"  # pas besoin de TreeTagger pour les tests
    for key, value in corpus_kwargs.items():
        setattr(config.corpus, key, value)
    return config


class TestDocxParser:
    """Tests pour le parseur DOCX"""

    def test_parse_latin_analyzer_docx(self, tmp_path):
        """Test parsing d'un DOCX latin_analyzer standard"""
        docx_path = tmp_path / "resultat.docx"
        text_lines = [
            "Dominus enim dicit in evangelio.",
            "Quod est ueritas.",
        ]
        _create_latin_analyzer_docx(docx_path, text_lines)

        config = _make_config(
            edition_id="Edi-7",
            title="Test Corpus",
            author="Augustinus",
        )
        parser = DocxParser(config)
        pages = parser.parse_file(docx_path)

        assert len(pages) == 1
        page = pages[0]
        assert len(page.lines) == 2
        assert page.lines[0] == "Dominus enim dicit in evangelio."
        assert page.lines[1] == "Quod est ueritas."

    def test_metadata_from_config(self, tmp_path):
        """Test que les metadonnees viennent du config.yaml"""
        docx_path = tmp_path / "test.docx"
        _create_latin_analyzer_docx(docx_path, ["Dominus dicit."])

        config = _make_config(
            edition_id="Edi-42",
            title="Summa Theologica",
            author="Thomas Aquinas",
            language="Latin",
            date="1274",
        )
        parser = DocxParser(config)
        pages = parser.parse_file(docx_path)

        page = pages[0]
        assert page.metadata.corpus_metadata["edition_id"] == "Edi-42"
        assert page.metadata.corpus_metadata["title"] == "Summa Theologica"
        assert page.metadata.corpus_metadata["author"] == "Thomas Aquinas"
        assert page.metadata.corpus_metadata["date"] == "1274"

    def test_folio_from_filename(self, tmp_path):
        """Test que le folio est derive du nom de fichier"""
        docx_path = tmp_path / "analyse_folio_42.docx"
        _create_latin_analyzer_docx(docx_path, ["Dominus dicit."])

        config = _make_config()
        parser = DocxParser(config)
        pages = parser.parse_file(docx_path)

        assert pages[0].metadata.folio == "analyse_folio_42"

    def test_skip_header(self, tmp_path):
        """Test que l'en-tete latin_analyzer est ignore"""
        docx_path = tmp_path / "test.docx"
        _create_latin_analyzer_docx(docx_path, ["Texte reel."])

        config = _make_config()
        parser = DocxParser(config)
        pages = parser.parse_file(docx_path)

        # Le texte ne doit pas contenir l'en-tete
        all_text = " ".join(pages[0].lines)
        assert "Analyse de texte latin" not in all_text
        assert "Légende" not in all_text
        assert "____" not in all_text
        assert "Texte reel." in all_text

    def test_parse_plain_docx(self, tmp_path):
        """Test parsing d'un DOCX sans en-tete latin_analyzer"""
        docx_path = tmp_path / "plain.docx"
        _create_plain_docx(docx_path, [
            "Dominus enim dicit.",
            "In evangelio scribitur.",
        ])

        config = _make_config()
        parser = DocxParser(config)
        pages = parser.parse_file(docx_path)

        assert len(pages) == 1
        assert len(pages[0].lines) == 2

    def test_parse_folder(self, tmp_path):
        """Test parsing d'un dossier de fichiers DOCX"""
        for i, text in enumerate(["Prima pars.", "Secunda pars."], start=1):
            path = tmp_path / f"page_{i}.docx"
            _create_latin_analyzer_docx(path, [text])

        config = _make_config(edition_id="Edi-1")
        parser = DocxParser(config)
        pages = parser.parse_folder(tmp_path)

        assert len(pages) == 2
        assert pages[0].metadata.page_number == 1
        assert pages[1].metadata.page_number == 2

    def test_empty_docx(self, tmp_path):
        """Test parsing d'un DOCX vide"""
        docx_path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(docx_path))

        config = _make_config()
        parser = DocxParser(config)
        pages = parser.parse_file(docx_path)

        assert len(pages) == 0

    def test_file_not_found(self, tmp_path):
        """Test erreur fichier non trouve"""
        config = _make_config()
        parser = DocxParser(config)

        with pytest.raises(FileNotFoundError):
            parser.parse_file(tmp_path / "inexistant.docx")


class TestDocxConverter:
    """Tests pour le convertisseur complet DOCX -> vertical"""

    def test_convert_file(self, tmp_path):
        """Test conversion complete fichier DOCX -> vertical"""
        from PAGEtopage.step5_docx import DocxConverter

        # Cree un DOCX
        docx_path = tmp_path / "test.docx"
        _create_latin_analyzer_docx(docx_path, [
            "Dominus enim dicit in evangelio.",
        ])

        # Configure
        config = _make_config(
            edition_id="Edi-7",
            title="Test",
            author="Anonymus",
            language="Latin",
        )

        # Convertit
        output_path = tmp_path / "corpus.vertical.txt"
        converter = DocxConverter(config)
        pages = converter.convert_file(docx_path, output_path)

        assert len(pages) == 1
        assert output_path.exists()

        # Verifie le contenu vertical
        content = output_path.read_text(encoding="utf-8")
        assert "<doc " in content
        assert "</doc>" in content
        assert "<s>" in content
        assert "</s>" in content
        assert 'edition_id="Edi-7"' in content

    def test_convert_folder(self, tmp_path):
        """Test conversion dossier DOCX -> vertical"""
        from PAGEtopage.step5_docx import DocxConverter

        # Cree des DOCX
        for i in range(3):
            path = tmp_path / "input" / f"page_{i}.docx"
            path.parent.mkdir(exist_ok=True)
            _create_latin_analyzer_docx(path, [f"Textus paginae {i}."])

        config = _make_config(edition_id="Edi-1")
        output_path = tmp_path / "corpus.vertical.txt"

        converter = DocxConverter(config)
        pages = converter.convert_folder(tmp_path / "input", output_path)

        assert len(pages) == 3
        assert output_path.exists()

    def test_vertical_has_tokens(self, tmp_path):
        """Test que le vertical contient des tokens lemmatises"""
        from PAGEtopage.step5_docx import DocxConverter

        docx_path = tmp_path / "test.docx"
        _create_latin_analyzer_docx(docx_path, ["Dominus dicit."])

        config = _make_config(edition_id="Edi-1")
        output_path = tmp_path / "corpus.vertical.txt"

        converter = DocxConverter(config)
        converter.convert_file(docx_path, output_path)

        content = output_path.read_text(encoding="utf-8")
        # Chaque token doit avoir 3 colonnes (mot, POS, lemme)
        lines = [l for l in content.split("\n") if l.strip() and not l.startswith("<")]
        for line in lines:
            parts = line.split("\t")
            assert len(parts) == 3, f"Ligne mal formatee: {line}"


class TestMultiPageDocx:
    """Tests pour le parsing de DOCX multi-pages (latin_analyzer v2.4+)"""

    def test_parse_multipage_docx(self, tmp_path):
        """Test parsing d'un DOCX avec plusieurs pages/folios"""
        docx_path = tmp_path / "multipage.docx"
        _create_multipage_docx(docx_path, [
            {
                "folio": "0001_r",
                "page_number": 1,
                "running_title": "De Civitate Dei",
                "lines": ["Dominus enim dicit.", "In evangelio scribitur."],
            },
            {
                "folio": "0001_v",
                "page_number": 2,
                "running_title": "De Civitate Dei",
                "lines": ["Quod est ueritas."],
            },
        ])

        config = _make_config(edition_id="Edi-7")
        parser = DocxParser(config)
        pages = parser.parse_file(docx_path)

        assert len(pages) == 2

        assert pages[0].metadata.folio == "0001_r"
        assert pages[0].metadata.page_number == 1
        assert pages[0].metadata.running_title == "De Civitate Dei"
        assert len(pages[0].lines) == 2
        assert pages[0].lines[0] == "Dominus enim dicit."

        assert pages[1].metadata.folio == "0001_v"
        assert pages[1].metadata.page_number == 2
        assert len(pages[1].lines) == 1
        assert pages[1].lines[0] == "Quod est ueritas."

    def test_parse_multipage_three_pages(self, tmp_path):
        """Test parsing avec 3 pages"""
        docx_path = tmp_path / "three_pages.docx"
        _create_multipage_docx(docx_path, [
            {"folio": "f1", "page_number": 1, "lines": ["Prima pars."]},
            {"folio": "f2", "page_number": 2, "lines": ["Secunda pars."]},
            {"folio": "f3", "page_number": 3, "lines": ["Tertia pars."]},
        ])

        config = _make_config()
        parser = DocxParser(config)
        pages = parser.parse_file(docx_path)

        assert len(pages) == 3
        assert pages[0].lines[0] == "Prima pars."
        assert pages[1].lines[0] == "Secunda pars."
        assert pages[2].lines[0] == "Tertia pars."

    def test_multipage_metadata_in_corpus(self, tmp_path):
        """Test que les metadonnees du corpus sont presentes dans chaque page"""
        docx_path = tmp_path / "meta.docx"
        _create_multipage_docx(docx_path, [
            {"folio": "f1", "page_number": 1, "lines": ["Texte 1."]},
            {"folio": "f2", "page_number": 2, "lines": ["Texte 2."]},
        ])

        config = _make_config(edition_id="Edi-42", author="Augustinus")
        parser = DocxParser(config)
        pages = parser.parse_file(docx_path)

        for page in pages:
            assert page.metadata.corpus_metadata["edition_id"] == "Edi-42"
            assert page.metadata.corpus_metadata["author"] == "Augustinus"

    def test_folio_only_header(self, tmp_path):
        """Test en-tete avec seulement le folio"""
        docx_path = tmp_path / "folio_only.docx"
        _create_multipage_docx(docx_path, [
            {"folio": "ms_42r", "lines": ["Dominus dicit."]},
        ])

        config = _make_config()
        parser = DocxParser(config)
        pages = parser.parse_file(docx_path)

        assert len(pages) == 1
        assert pages[0].metadata.folio == "ms_42r"

    def test_page_header_not_in_text(self, tmp_path):
        """Test que les en-tetes de page ne sont pas dans le texte"""
        docx_path = tmp_path / "no_header_in_text.docx"
        _create_multipage_docx(docx_path, [
            {"folio": "f1", "page_number": 1, "lines": ["Texte reel."]},
        ])

        config = _make_config()
        parser = DocxParser(config)
        pages = parser.parse_file(docx_path)

        all_text = " ".join(pages[0].lines)
        assert "Folio:" not in all_text
        assert "─" not in all_text
        assert "Texte reel." in all_text


class TestHeaderDetection:
    """Tests pour la detection de l'en-tete"""

    def test_header_markers(self):
        """Test que les marqueurs d'en-tete sont bien definis"""
        assert any("Analyse" in m for m in HEADER_MARKERS)
        assert any("Légende" in m for m in HEADER_MARKERS)
        assert any("Noir" in m for m in HEADER_MARKERS)

    def test_is_header_content(self):
        """Test detection contenu en-tete"""
        assert DocxParser._is_header_content("Analyse de texte latin medieval")
        assert DocxParser._is_header_content("Légende : ...")
        assert DocxParser._is_header_content("Noir = OK (score ≥75)")
        assert not DocxParser._is_header_content("Dominus dicit in evangelio.")

    def test_page_header_pattern(self):
        """Test detection des en-tetes de page/folio"""
        # Format complet
        text = "──────────────────── Folio: 0001_r | Page: 1 | De Civitate Dei ────────────────────"
        info = _parse_page_header(text)
        assert info is not None
        assert info["folio"] == "0001_r"
        assert info["page_number"] == 1
        assert info["running_title"] == "De Civitate Dei"

    def test_page_header_folio_only(self):
        """Test detection en-tete avec folio seul"""
        text = "──────────────────── Folio: ms_42r ────────────────────"
        info = _parse_page_header(text)
        assert info is not None
        assert info["folio"] == "ms_42r"

    def test_page_header_no_match_normal_text(self):
        """Test que du texte normal ne matche pas"""
        assert _parse_page_header("Dominus dicit.") is None
        assert _parse_page_header("Analyse de texte latin") is None
