#!/usr/bin/env python3
"""
Analyseur de textes latins médiévaux - Version 2.1
================================================

Système intégré utilisant :
- PyCollatinus pour l'analyse morphologique
- Dictionnaire Du Cange (99k+ entrées de latin médiéval)
- Système de scoring multi-critères
- Colorisation à 3 niveaux (rouge/orange/noir)
- Support XML Pages avec extraction MainZone automatique
- Normalisation u/v et i/j
- Fusion des mots avec tirets
- Filtrage des chiffres romains

Auteur: Claude
Date: 2025-11-25
"""

import argparse
import os
import re
import sys
import docx
from docx.shared import RGBColor, Cm, Pt
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path
from collections import Counter
from typing import List, Tuple, Optional

# Ajouter PyCollatinus au path
sys.path.insert(0, '/tmp/collatinus-python')

from pycollatinus import Lemmatiseur

# Import local si exécuté comme script, sinon import relatif
try:
    from page_xml_parser import PageXMLParser
except ImportError:
    from .page_xml_parser import PageXMLParser


# Patterns regex compilés
PATTERNS = {
    # Pattern pour capturer le dernier mot avec tiret en fin de ligne
    # Non-greedy pour éviter de capturer tout le texte avant
    'hyphenated_word': re.compile(r'^(.*\s)?(\w+)-\s*$'),
    # Pattern pour chiffres romains avec point optionnel (? = 0 ou 1)
    'roman_numeral': re.compile(r'^[ivxlcdm]+\.?$'),
}


class LatinAnalyzer:
    """
    Analyseur de textes latins médiévaux avec scoring multi-critères

    Utilise deux sources lexicales complémentaires pour valider chaque mot :
    - PyCollatinus : ~500 000 formes de latin classique (analyse morphologique)
    - Dictionnaire Du Cange : ~99 000 entrées de latin médiéval

    Algorithme de scoring (0-100 points, base = 50) :
        - Critère 1 : Reconnu par PyCollatinus (latin classique) → +30 pts
        - Critère 2 : Présent dans le dictionnaire Du Cange (médiéval) → +40 pts
        - Critère 3 : Suffixe médiéval productif (-arius, -atio, -torium...) → +10 pts
        - Critère 4 : Contexte ecclésiastique (mots voisins reconnus) → +5 pts
        - Critère 5 : Variante orthographique médiévale (ae↔e, ti↔ci) → +10 pts

    Seuils de colorisation :
        - Score >= 75 → Noir (mot validé)
        - Score 40-74 → Orange (à vérifier)
        - Score < 40  → Rouge (erreur probable)

    Pré-traitements appliqués :
        - Normalisation u/v et i/j (variantes médiévales)
        - Filtrage des chiffres romains avec point (xuiii., uii., etc.)
        - Fusion des mots coupés par tiret en fin de ligne

    Usage:
        analyzer = LatinAnalyzer(ducange_dict_file="dictionnaire_ducange.txt")
        results = analyzer.analyze_text_file("mon_texte.txt")
        analyzer.generate_docx("resultat.docx", results)
    """

    def __init__(self, ducange_dict_file=None):
        """
        Initialise l'analyseur.

        Args:
            ducange_dict_file (str): Chemin vers le dictionnaire Du Cange
        """
        print("🔄 Initialisation de l'analyseur latin...")

        # Charger PyCollatinus
        print("  📚 Chargement de Collatinus (latin classique)...")
        self.lemmatizer = Lemmatiseur()
        print("  ✅ Collatinus chargé")

        # Charger le dictionnaire médiéval
        self.medieval_dict = set()

        # Compteurs de debug
        self.ducange_matches_count = 0
        self.ducange_examples = []
        self.collatinus_matches_count = 0
        self.collatinus_examples = []
        self.both_matches_count = 0

        if ducange_dict_file and os.path.exists(ducange_dict_file):
            print(f"  📚 Chargement du dictionnaire Du Cange...")
            print(f"  📁 Fichier : {ducange_dict_file}")
            with open(ducange_dict_file, 'r', encoding='utf-8') as f:
                for line in f:
                    word = line.strip().lower()
                    if word:
                        # Ajouter aussi les variantes normalisées
                        self.medieval_dict.add(word)
                        self.medieval_dict.add(self.normalize_word(word))
            print(f"  ✅ {len(self.medieval_dict)} mots médiévaux chargés")
            print(f"  🔍 Exemples de mots Du Cange : {list(self.medieval_dict)[:10]}")
        else:
            print("  ⚠️  Dictionnaire Du Cange non trouvé")
            if ducange_dict_file:
                print(f"  ⚠️  Chemin testé : {ducange_dict_file}")

        # Suffixes typiques du latin médiéval
        self.medieval_suffixes = [
            'arius', 'aria', 'arium',
            'atio', 'ationis',
            'tor', 'toris',
            'torium', 'torii',
            'mentum', 'menti',
            'itia', 'itiae',
        ]

        # Contexte ecclésiastique
        self.ecclesiastical_words = {
            'abbas', 'abbatia', 'abbatissa', 'archiepiscopus', 'basilica',
            'canonicus', 'capitulum', 'cardinalis', 'clericus', 'diaconus',
            'diocesis', 'dominus', 'ecclesia', 'episcopus', 'monasterium',
            'monachus', 'parochia', 'presbyter', 'sacerdos', 'sanctus',
        }

        print("✅ Analyseur prêt\n")

    @staticmethod
    def normalize_word(word: str) -> str:
        """
        Normalise un mot latin en convertissant u→v et i→j.

        Les variantes médiévales u/v et i/j sont considérées comme identiques :
        - uel → vel
        - uidetur → videtur
        - iustus → iustus (j→i déjà fait en médiéval)

        Args:
            word (str): Mot à normaliser

        Returns:
            str: Mot normalisé
        """
        word = word.replace('u', 'v').replace('U', 'V')
        word = word.replace('j', 'i').replace('J', 'I')
        return word

    @staticmethod
    def is_roman_numeral_with_dot(word: str) -> bool:
        """
        Détecte si un mot est un chiffre romain avec point final.

        Gère les variantes médiévales avec u au lieu de v :
        - xuiii. (14)
        - uii. (7)
        - ui. (6)

        Args:
            word (str): Mot à vérifier

        Returns:
            bool: True si c'est un chiffre romain avec point
        """
        # Normaliser u→v pour détecter les chiffres romains médiévaux
        normalized = word.lower().replace('u', 'v')
        return bool(PATTERNS['roman_numeral'].match(normalized))

    def merge_hyphenated_words(self, lines: List[str]) -> List[str]:
        """
        Fusionne les mots coupés avec trait d'union entre deux lignes.

        Inspiré de xml_corpus_processor._merge_hyphenated_words()

        Args:
            lines: Liste des lignes à traiter

        Returns:
            Liste des lignes avec mots fusionnés

        Example:
            >>> lines = ["sancti-", "tatis"]
            >>> analyzer.merge_hyphenated_words(lines)
            ["sanctitatis"]
        """
        processed_lines = []
        fusion_count = 0
        i = 0

        while i < len(lines):
            current_line = lines[i].strip() if isinstance(lines[i], str) else ''

            # Détection des mots coupés avec trait d'union
            match = PATTERNS['hyphenated_word'].search(current_line)

            if match and i + 1 < len(lines):
                prefix_text = match.group(1) if match.group(1) else ''  # Texte avant le mot coupé
                prefix_word = match.group(2)  # Première partie du mot coupé

                # Récupération de la ligne suivante
                next_line = lines[i + 1].strip() if isinstance(lines[i + 1], str) else ''
                next_line_parts = next_line.split(' ', 1)

                if next_line_parts:
                    # Fusion du mot
                    suffix_word = next_line_parts[0]
                    fused_word = prefix_word + suffix_word

                    # Construction de la nouvelle ligne
                    new_line = prefix_text + fused_word

                    # Ajout du reste de la ligne suivante si présent
                    if len(next_line_parts) > 1:
                        new_line += ' ' + next_line_parts[1]

                    # Debug : afficher les 5 premières fusions
                    fusion_count += 1
                    if fusion_count <= 5:
                        print(f"  🔗 Fusion #{fusion_count}: '{prefix_word}-' + '{suffix_word}' → '{fused_word}'")

                    processed_lines.append(new_line)
                    i += 2  # Sauter la ligne suivante
                    continue

            # Pas de fusion nécessaire
            processed_lines.append(current_line)
            i += 1

        if fusion_count > 0:
            print(f"  ✅ {fusion_count} mots fusionnés au total")
        else:
            print(f"  ⚠️  Aucun mot avec tiret détecté pour fusion")

        return processed_lines

    def analyze_word(self, word, context_words=None):
        """
        Analyse un mot avec scoring multi-critères.

        Args:
            word (str): Le mot à analyser
            context_words (list): Liste des mots environnants

        Returns:
            dict: {
                'word': str,
                'recognized_classical': bool,
                'recognized_medieval': bool,
                'confidence_score': int (0-100),
                'reasons': list,
                'color_code': str ('black', 'orange', 'red')
            }
        """
        clean_word = word.lower().strip()
        context_words = context_words or []

        result = {
            'word': word,
            'recognized_classical': False,
            'recognized_medieval': False,
            'confidence_score': 50,  # Score de base neutre
            'reasons': [],
            'color_code': 'orange'
        }

        # Ignorer les mots très courts
        if len(clean_word) < 2:
            result['confidence_score'] = 100
            result['color_code'] = 'black'
            result['reasons'].append("mot trop court pour être analysé")
            return result

        # Ignorer les chiffres romains avec point (xuiii., uii., ui., etc.)
        if self.is_roman_numeral_with_dot(clean_word):
            result['confidence_score'] = 100
            result['color_code'] = 'black'
            result['reasons'].append("chiffre romain")
            return result

        # Normaliser pour la comparaison (u→v, i→j)
        normalized_word = self.normalize_word(clean_word)

        # Critère 1 : Reconnu par Collatinus (latin classique)
        try:
            # Tester le mot original ET la version normalisée
            for test_word in [clean_word, normalized_word]:
                # CORRECTION : Convertir generator en liste
                analyses = list(self.lemmatizer.lemmatise(test_word))
                if analyses and len(analyses) > 0:
                    result['recognized_classical'] = True
                    result['confidence_score'] += 30
                    result['reasons'].append(f"latin classique valide ({len(analyses)} analyse(s))")

                    # Debug : compter
                    self.collatinus_matches_count += 1
                    if len(self.collatinus_examples) < 20:
                        self.collatinus_examples.append(clean_word)
                    break
        except Exception as e:
            # Debug : afficher les erreurs PyCollatinus pour diagnostic
            pass

        # Critère 2 : Présent dans le dictionnaire Du Cange (avec normalisation)
        if clean_word in self.medieval_dict or normalized_word in self.medieval_dict:
            result['recognized_medieval'] = True
            result['confidence_score'] += 40
            result['reasons'].append("présent dans le dictionnaire Du Cange")

            # Debug : compter et enregistrer des exemples
            self.ducange_matches_count += 1
            if len(self.ducange_examples) < 20:
                self.ducange_examples.append(clean_word)

            # Compter si reconnu par les deux
            if result['recognized_classical']:
                self.both_matches_count += 1

        # Critère 3 : Suffixe médiéval typique
        for suffix in self.medieval_suffixes:
            if normalized_word.endswith(suffix):
                result['confidence_score'] += 10
                result['reasons'].append(f"suffixe médiéval productif (-{suffix})")
                break

        # Critère 4 : Contexte ecclésiastique
        context_ecclesiastical = any(
            self.normalize_word(w.lower()) in self.ecclesiastical_words
            for w in context_words
        )
        if context_ecclesiastical:
            result['confidence_score'] += 5
            result['reasons'].append("contexte ecclésiastique")

        # Critère 5 : Variante orthographique médiévale (ae→e, ti→ci, etc.)
        if self._is_medieval_variant(normalized_word):
            result['confidence_score'] += 10
            result['reasons'].append("variante orthographique médiévale détectée")

        # Déterminer la couleur selon le score
        score = min(result['confidence_score'], 100)
        if score >= 75:
            result['color_code'] = 'black'
        elif score >= 40:
            result['color_code'] = 'orange'
        else:
            result['color_code'] = 'red'

        result['confidence_score'] = score
        return result

    def _is_medieval_variant(self, word):
        """
        Détecte si un mot pourrait être une variante orthographique médiévale.

        Args:
            word (str): Le mot à analyser (déjà normalisé)

        Returns:
            bool: True si variante détectée
        """
        # Générer des variantes classiques et vérifier si elles sont reconnues
        variants = []

        # ae ← e (medieval → mediaeval)
        if 'e' in word and 'ae' not in word:
            variants.append(word.replace('e', 'ae', 1))

        # ti ← ci (gracia → gratia)
        if 'ci' in word:
            variants.append(word.replace('ci', 'ti'))

        # Vérifier si une variante existe dans les dictionnaires
        for variant in variants:
            if variant in self.medieval_dict:
                return True
            try:
                analyses = self.lemmatizer.lemmatise(variant)
                if analyses and len(analyses) > 0:
                    return True
            except Exception:
                pass

        return False

    def extract_and_process_xml(self, xml_path, column_mode='single'):
        """
        Extrait le texte depuis XML Pages et traite les tirets.

        Args:
            xml_path (str): Chemin vers fichier XML ou dossier
            column_mode (str): 'single' ou 'dual'

        Returns:
            List[str]: Lignes de texte traitées
        """
        print(f"📄 Extraction du texte depuis XML Pages...")

        parser = PageXMLParser(column_mode=column_mode)

        # Extraire le texte
        if os.path.isfile(xml_path):
            lines, metadata = parser.parse_file(xml_path)
            print(f"  ✅ 1 fichier traité, {len(lines)} lignes extraites")
        elif os.path.isdir(xml_path):
            text, metadata_list = parser.parse_folder(xml_path)
            lines = text.split('\n')
            print(f"  ✅ {len(metadata_list)} fichiers traités, {len(lines)} lignes extraites")
        else:
            raise ValueError(f"Chemin invalide : {xml_path}")

        # Fusionner les mots avec tirets
        print(f"🔗 Fusion des mots avec tirets...")
        lines = self.merge_hyphenated_words(lines)
        print(f"  ✅ Fusion terminée")

        return lines

    def analyze_page_xml(self, input_path, column_mode='single'):
        """
        Analyse un fichier ou dossier XML Pages.

        Args:
            input_path (str): Chemin vers fichier XML ou dossier
            column_mode (str): 'single' ou 'dual'

        Returns:
            dict: Statistiques et résultats de l'analyse
        """
        # Extraire et traiter le texte XML
        lines = self.extract_and_process_xml(input_path, column_mode)

        # Analyser les lignes
        return self._analyze_lines(lines, source=input_path)

    def analyze_text_file(self, input_file):
        """
        Analyse un fichier texte brut complet.

        Args:
            input_file (str): Chemin vers le fichier texte

        Returns:
            dict: Statistiques et résultats de l'analyse
        """
        print(f"📄 Analyse du fichier texte : {input_file}")

        with open(input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        # Fusionner les mots avec tirets
        print(f"🔗 Fusion des mots avec tirets...")
        lines = self.merge_hyphenated_words(lines)

        return self._analyze_lines(lines, source=input_file)

    def _analyze_lines(self, lines, source):
        """
        Analyse une liste de lignes de texte.

        Args:
            lines (list): Liste des lignes à analyser
            source (str): Source du texte (pour logs)

        Returns:
            dict: Statistiques et résultats de l'analyse
        """
        print(f"🔍 Analyse en cours...")
        print(f"   Source : {source}")

        results = []
        word_counter = Counter()
        score_distribution = {'black': 0, 'orange': 0, 'red': 0}

        total_words = 0

        for line_num, line in enumerate(lines, 1):
            # Extraire les mots de la ligne
            words = re.findall(r'\b[a-zA-ZÀ-ÿ]+\b', line)

            for i, word in enumerate(words):
                if len(word) < 2:
                    continue

                total_words += 1

                # Contexte = mots environnants (±3 mots)
                context_start = max(0, i - 3)
                context_end = min(len(words), i + 4)
                context = words[context_start:i] + words[i+1:context_end]

                # Analyser le mot
                analysis = self.analyze_word(word, context)

                results.append({
                    'line': line_num,
                    'word': word,
                    'analysis': analysis
                })

                word_counter[word.lower()] += 1
                score_distribution[analysis['color_code']] += 1

                # Afficher progression tous les 1000 mots
                if total_words % 1000 == 0:
                    print(f"  ⏳ {total_words} mots analysés...")

        if total_words == 0:
            print("⚠️  Aucun mot trouvé à analyser")
            return {
                'results': [],
                'statistics': {
                    'total_words': 0,
                    'distribution': score_distribution,
                    'word_frequency': word_counter
                },
                'source_lines': lines
            }

        print(f"✅ Analyse terminée : {total_words} mots traités")
        print(f"\n📊 Distribution des scores :")
        print(f"  ✅ Noir (bons mots)      : {score_distribution['black']} ({score_distribution['black']*100//total_words}%)")
        print(f"  ⚠️  Orange (douteux)      : {score_distribution['orange']} ({score_distribution['orange']*100//total_words}%)")
        print(f"  ❌ Rouge (erreurs prob.) : {score_distribution['red']} ({score_distribution['red']*100//total_words}%)")

        # Statistiques détaillées par source
        print(f"\n📚 Statistiques de reconnaissance par source :")

        if self.collatinus_matches_count > 0:
            print(f"  🏛️  PyCollatinus (latin classique) : {self.collatinus_matches_count} mots")
            print(f"      Exemples : {', '.join(self.collatinus_examples[:10])}")
        else:
            print(f"  ⚠️  PyCollatinus : 0 mots reconnus (vérifier l'installation)")

        if self.ducange_matches_count > 0:
            print(f"  📖 Du Cange (latin médiéval) : {self.ducange_matches_count} mots")
            print(f"      Exemples : {', '.join(self.ducange_examples[:10])}")
        else:
            print(f"  ⚠️  Du Cange : 0 mots reconnus (vérifier le dictionnaire)")

        if self.both_matches_count > 0:
            print(f"  🔗 Reconnus par les deux : {self.both_matches_count} mots")

        # Calculer les mots reconnus uniquement par l'une des sources
        only_collatinus = self.collatinus_matches_count - self.both_matches_count
        only_ducange = self.ducange_matches_count - self.both_matches_count

        print(f"\n  📊 Répartition :")
        print(f"      Uniquement PyCollatinus : {only_collatinus}")
        print(f"      Uniquement Du Cange : {only_ducange}")
        print(f"      Les deux : {self.both_matches_count}")

        return {
            'results': results,
            'statistics': {
                'total_words': total_words,
                'distribution': score_distribution,
                'word_frequency': word_counter
            },
            'source_lines': lines  # Garder les lignes sources pour le DOCX
        }

    def generate_docx(self, output_docx, analysis_results):
        """
        Génère un document Word avec colorisation à 3 niveaux.

        Args:
            output_docx (str): Fichier DOCX de sortie
            analysis_results (dict): Résultats de l'analyse
        """
        print(f"\n📝 Génération du document Word...")

        # Récupérer les lignes sources
        original_lines = analysis_results.get('source_lines', [])

        if not original_lines:
            print("❌ Erreur : Pas de lignes sources disponibles")
            return

        # Créer un index des analyses par ligne et mot
        analysis_index = {}
        for item in analysis_results['results']:
            line_num = item['line']
            word = item['word'].lower()
            if line_num not in analysis_index:
                analysis_index[line_num] = {}
            analysis_index[line_num][word] = item['analysis']

        # Créer le document
        doc = docx.Document()

        # Configuration de la page
        section = doc.sections[0]
        section.page_width = Cm(40)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        # Ajouter un titre
        title = doc.add_paragraph()
        title_run = title.add_run("Analyse de texte latin médiéval")
        title_run.bold = True
        title_run.font.size = Pt(14)

        # Légende
        legend = doc.add_paragraph()
        legend.add_run("Légende : ").bold = True

        black_run = legend.add_run("Noir = OK (score ≥75) ")
        black_run.font.color.rgb = RGBColor(0, 0, 0)

        orange_run = legend.add_run("Orange = À vérifier (score 40-74) ")
        orange_run.font.color.rgb = RGBColor(255, 140, 0)

        red_run = legend.add_run("Rouge = Erreur probable (score <40)")
        red_run.font.color.rgb = RGBColor(255, 0, 0)

        doc.add_paragraph("_" * 80)

        # Traiter chaque ligne
        for line_num, line in enumerate(original_lines, 1):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            # Découper en mots et espaces
            tokens = re.findall(r'(\S+|\s+)', line)

            for token in tokens:
                if token.isspace():
                    paragraph.add_run(token)
                else:
                    # Extraire le mot pur (sans ponctuation)
                    clean_token = ''.join(c for c in token if c.isalpha()).lower()

                    # Récupérer l'analyse
                    color = RGBColor(0, 0, 0)  # Noir par défaut
                    if line_num in analysis_index and clean_token in analysis_index[line_num]:
                        analysis = analysis_index[line_num][clean_token]
                        if analysis['color_code'] == 'red':
                            color = RGBColor(255, 0, 0)
                        elif analysis['color_code'] == 'orange':
                            color = RGBColor(255, 140, 0)

                    run = paragraph.add_run(token)
                    run.font.color.rgb = color
                    run.font.size = Pt(10)

        # Sauvegarder
        doc.save(output_docx)
        print(f"✅ Document créé : {output_docx}")


def analyze_orange_patterns(words_list):
    """
    Analyse les patterns des mots orange (non reconnus).

    Args:
        words_list (list): Liste de mots

    Returns:
        dict: Statistiques et patterns détectés
    """
    if not words_list:
        return None

    stats = {
        'total': len(words_list),
        'unique': len(set(words_list)),
        'avg_length': sum(len(w) for w in words_list) / len(words_list),
        'length_distribution': Counter(len(w) for w in words_list),
        'frequency': Counter(words_list),
        'patterns': {
            'contains_ae': sum(1 for w in words_list if 'ae' in w),
            'contains_oe': sum(1 for w in words_list if 'oe' in w),
            'contains_ph': sum(1 for w in words_list if 'ph' in w),
            'double_consonants': sum(1 for w in words_list if re.search(r'([bcdfghjklmnpqrstvwxz])\1', w)),
            'ends_with_us': sum(1 for w in words_list if w.endswith('us')),
            'ends_with_is': sum(1 for w in words_list if w.endswith('is')),
            'ends_with_um': sum(1 for w in words_list if w.endswith('um')),
            'ends_with_a': sum(1 for w in words_list if w.endswith('a')),
            'starts_with_vowel': sum(1 for w in words_list if w[0] in 'aeiou'),
            'very_short': sum(1 for w in words_list if len(w) <= 3),
            'very_long': sum(1 for w in words_list if len(w) >= 12),
        }
    }

    return stats


def categorize_orange_words(words_list):
    """
    Catégorise les mots orange par type potentiel.

    Args:
        words_list (list): Liste de mots

    Returns:
        dict: Mots catégorisés
    """
    categories = {
        'possible_abbreviations': [],
        'possible_ocr_errors': [],
        'medieval_spelling': [],
    }

    for word in set(words_list):
        # Abréviations potentielles
        if len(word) <= 3:
            categories['possible_abbreviations'].append(word)

        # Erreurs OCR potentielles
        if re.search(r'[0-9]|rn|cl(?!a)|li(?![a-z])', word):
            categories['possible_ocr_errors'].append(word)

        # Variantes orthographiques médiévales
        if 'ae' in word or 'oe' in word or 'ph' in word or re.search(r'([bcdfghjklmnpqrstvwxz])\1', word):
            categories['medieval_spelling'].append(word)

    return categories


def generate_orange_report(analysis_results, report_file):
    """
    Génère un rapport détaillé des mots orange dans un fichier.

    Args:
        analysis_results (dict): Résultats de l'analyse
        report_file (str): Chemin du fichier de rapport
    """
    # Extraire les mots orange
    orange_words = []
    for item in analysis_results['results']:
        if item['analysis']['color_code'] == 'orange':
            orange_words.append(item['word'].lower())

    if not orange_words:
        print("  ✅ Aucun mot orange - pas de rapport généré")
        return

    print(f"\n📝 Génération du rapport d'analyse des mots orange...")

    # Analyser les patterns
    stats = analyze_orange_patterns(orange_words)
    categories = categorize_orange_words(orange_words)

    # Écrire le rapport
    with open(report_file, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("RAPPORT D'ANALYSE DES MOTS ORANGE (NON RECONNUS)\n")
        f.write("=" * 80 + "\n\n")

        # Statistiques générales
        f.write(" 📊 STATISTIQUES GÉNÉRALES ".center(80, '-') + "\n")
        f.write(f"  Total de mots orange : {stats['total']}\n")
        f.write(f"  Mots uniques : {stats['unique']}\n")
        f.write(f"  Longueur moyenne : {stats['avg_length']:.1f} caractères\n\n")

        # Distribution par longueur
        f.write(" 📏 DISTRIBUTION PAR LONGUEUR ".center(80, '-') + "\n")
        for length in sorted(stats['length_distribution'].keys())[:15]:
            count = stats['length_distribution'][length]
            percent = count * 100 / stats['total']
            bar = '█' * int(percent / 2)
            f.write(f"  {length:2d} caractères : {count:4d} ({percent:5.1f}%) {bar}\n")
        f.write("\n")

        # Patterns détectés
        f.write(" 🔍 PATTERNS DÉTECTÉS ".center(80, '-') + "\n")
        for pattern, count in stats['patterns'].items():
            if count > 0:
                percent = count * 100 / stats['total']
                f.write(f"  {pattern:25s} : {count:4d} ({percent:5.1f}%)\n")
        f.write("\n")

        # TOP 50 mots les plus fréquents
        f.write(" 🏆 TOP 50 MOTS LES PLUS FRÉQUENTS ".center(80, '-') + "\n")
        f.write(f"  {'Rang':<6} {'Fréquence':<12} {'Mot'}\n")
        f.write(f"  {'-'*6} {'-'*12} {'-'*50}\n")
        for rank, (word, count) in enumerate(stats['frequency'].most_common(50), 1):
            f.write(f"  {rank:<6d} {count:<12d} {word}\n")
        f.write("\n")

        # Catégories
        f.write(" 📂 CATÉGORISATION DES MOTS ".center(80, '-') + "\n\n")
        for category, words in categories.items():
            if words:
                f.write(f"  {category.replace('_', ' ').title()} ({len(words)} mots) :\n")
                for word in sorted(words)[:15]:
                    freq = stats['frequency'][word.lower()]
                    f.write(f"    - {word} (x{freq})\n")
                if len(words) > 15:
                    f.write(f"    ... et {len(words) - 15} autres\n")
                f.write("\n")

        # Recommandations
        f.write(" 💡 RECOMMANDATIONS ".center(80, '-') + "\n\n")

        if stats['patterns']['contains_ae'] > 10:
            f.write("  ✓ Ajouter la normalisation ae ↔ e (variantes médiévales)\n")

        if stats['patterns']['contains_ph'] > 5:
            f.write("  ✓ Ajouter la normalisation ph ↔ f\n")

        if stats['patterns']['double_consonants'] > 20:
            f.write("  ✓ Tester les variantes avec géminées simplifiées (mm→m, nn→n)\n")

        if len(categories['possible_abbreviations']) > 20:
            f.write("  ✓ Créer un dictionnaire d'abréviations médiévales courantes\n")

        if len(categories['possible_ocr_errors']) > 30:
            f.write("  ✓ Implémenter la correction d'erreurs OCR (rn→m, cl→d, etc.)\n")

        f.write("\n")

        # Estimations d'impact
        f.write(" 📈 ESTIMATION D'IMPACT DES AMÉLIORATIONS ".center(80, '-') + "\n\n")

        # Variantes orthographiques
        variants_impact = (stats['patterns']['contains_ae'] +
                          stats['patterns']['contains_ph'] +
                          stats['patterns']['double_consonants'])
        variants_percent = variants_impact * 100 / stats['total']
        f.write(f"  Variantes orthographiques : +{variants_percent:.1f}% potentiel\n")

        # Abréviations
        abbrev_percent = len(categories['possible_abbreviations']) * 100 / stats['total']
        f.write(f"  Dictionnaire abréviations : +{abbrev_percent:.1f}% potentiel\n")

        # Total estimé
        total_stats = analysis_results['statistics']
        current_percent = total_stats['distribution']['black'] * 100 / total_stats['total_words']
        total_improvement = min(variants_percent + abbrev_percent,
                               total_stats['distribution']['orange'] * 100 / total_stats['total_words'])
        f.write(f"\n  🎯 AMÉLIORATION TOTALE ESTIMÉE : +{total_improvement:.1f}% ")
        f.write(f"({current_percent:.0f}% → {current_percent + total_improvement:.0f}%)\n\n")

        f.write("=" * 80 + "\n")

    print(f"  ✅ Rapport généré : {report_file}")


def main():
    """Fonction principale avec arguments CLI."""
    parser = argparse.ArgumentParser(
        description='Analyseur de textes latins médiévaux avec détection intelligente des erreurs',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemples d'utilisation :

  # Analyser un fichier texte
  python3 latin_analyzer_v2.py -i texte.txt -o resultat.docx

  # Analyser des XML Pages (mode single)
  python3 latin_analyzer_v2.py -i corpus_xml/ -o resultat.docx -m xml-single

  # Analyser des XML Pages (mode dual - 2 colonnes)
  python3 latin_analyzer_v2.py -i corpus_xml/ -o resultat.docx -m xml-dual

  # Avec rapport détaillé des mots orange (non reconnus)
  python3 latin_analyzer_v2.py -i texte.txt -o resultat.docx --report analyse_orange.txt
        """
    )

    parser.add_argument(
        '-i', '--input',
        required=True,
        help='Fichier texte ou dossier XML Pages à analyser'
    )

    parser.add_argument(
        '-o', '--output',
        required=True,
        help='Fichier DOCX de sortie'
    )

    parser.add_argument(
        '-d', '--ducange',
        default=None,
        help='Chemin vers le dictionnaire Du Cange (optionnel)'
    )

    parser.add_argument(
        '-m', '--mode',
        choices=['txt', 'xml-single', 'xml-dual'],
        default='txt',
        help='Mode d\'analyse : txt (fichier texte), xml-single (XML 1 colonne), xml-dual (XML 2 colonnes)'
    )

    parser.add_argument(
        '--report',
        default=None,
        help='Générer un rapport détaillé des mots orange (non reconnus) dans le fichier spécifié'
    )

    args = parser.parse_args()

    # Vérifier que le fichier d'entrée existe
    if not os.path.exists(args.input):
        print(f"❌ Erreur : Le fichier/dossier d'entrée n'existe pas : {args.input}")
        sys.exit(1)

    # Déterminer le chemin du dictionnaire par défaut si non fourni
    if args.ducange is None:
        script_dir = Path(__file__).parent.parent
        default_ducange = script_dir / "data" / "ducange_data" / "dictionnaire_ducange.txt"
        if default_ducange.exists():
            args.ducange = str(default_ducange)
            print(f"💡 Utilisation du dictionnaire par défaut : {args.ducange}\n")
        else:
            print("⚠️  Pas de dictionnaire Du Cange spécifié (analyse latin classique uniquement)\n")

    print("=" * 70)
    print("  ANALYSEUR DE TEXTES LATINS MÉDIÉVAUX - VERSION 2.1")
    print("  PyCollatinus + Du Cange + Scoring Multi-critères")
    print("=" * 70)
    print()

    # Initialiser l'analyseur
    analyzer = LatinAnalyzer(ducange_dict_file=args.ducange)

    # Analyser selon le mode
    if args.mode == 'txt':
        analysis_results = analyzer.analyze_text_file(args.input)
    elif args.mode == 'xml-single':
        analysis_results = analyzer.analyze_page_xml(args.input, column_mode='single')
    elif args.mode == 'xml-dual':
        analysis_results = analyzer.analyze_page_xml(args.input, column_mode='dual')

    # Générer le document Word
    analyzer.generate_docx(args.output, analysis_results)

    # Générer le rapport d'analyse des mots orange si demandé
    if args.report:
        generate_orange_report(analysis_results, args.report)

    print("\n" + "=" * 70)
    print("✅ TRAITEMENT TERMINÉ !")
    print(f"📁 Fichier généré : {args.output}")
    if args.report:
        print(f"📊 Rapport orange généré : {args.report}")
    print("=" * 70)

    return 0


if __name__ == "__main__":
    sys.exit(main())
