#!/usr/bin/env python3
"""
Export Nakala - Version par ŒUVRE (séparée)

Logique : Un dossier = un fichier vertical (= une œuvre)
Si plusieurs œuvres partagent le même ID, elles ont chacune leur propre dossier.
La fiche est dupliquée dans chaque dossier concerné.

Résultat attendu : 122 dossiers (1 par fichier vertical)

Structure de sortie :
    Export_Nakala/
    ├── Libre_de_droits/
    │   ├── Oeuvre_A_Edi-41/
    │   │   ├── fiche.docx      (même fiche, dupliquée)
    │   │   ├── vertical.txt
    │   │   └── textes/
    │   └── Oeuvre_B_Edi-41/
    │       ├── fiche.docx      (même fiche, dupliquée)
    │       ├── vertical.txt
    │       └── textes/
    └── Non_libre_de_droits/
        └── ...

Auteur: Script pour le projet CiSaMe
"""

import os
import re
import glob
import shutil
from datetime import datetime
from collections import defaultdict
from docx import Document


def normalize_filename(name, max_length=80):
    """Normalise un nom pour l'utiliser comme nom de dossier."""
    if not name:
        return "Inconnu"
    # Remplacer les caractères problématiques
    name = re.sub(r'[<>:"/\\|?*]', '_', name)
    name = re.sub(r'\s+', '_', name)
    name = re.sub(r'_+', '_', name)
    name = name.strip('_')
    # Limiter la longueur
    if len(name) > max_length:
        name = name[:max_length].rstrip('_')
    return name


def extract_info_from_docx(filepath):
    """
    Extrait l'identifiant édition et le statut libre de droits d'une fiche .docx.
    
    Returns:
        dict: {'id': 'Edi-XX' ou None, 'titre_oeuvre': str, 'libre_de_droits': bool}
    """
    result = {'id': None, 'titre_oeuvre': None, 'libre_de_droits': False}
    
    try:
        doc = Document(filepath)
        text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Chercher l'identifiant édition
        match = re.search(r'Identifiant\s+[ée]dition\s*:\s*(Edi-\d+)', text, re.IGNORECASE)
        if match:
            result['id'] = match.group(1)
        
        # Chercher le titre de l'œuvre
        titre_match = re.search(r'Titre\s*:\s*\*?([^*\n]+)\*?', text)
        if titre_match:
            result['titre_oeuvre'] = titre_match.group(1).strip()
        
        # Chercher si libre de droits
        if re.search(r'Libre\s+de\s+droits\s*:\s*Oui', text, re.IGNORECASE):
            result['libre_de_droits'] = True
    except Exception as e:
        print(f"  Erreur lecture {filepath}: {e}")
    
    return result


def extract_info_from_vertical(filepath):
    """
    Extrait l'identifiant édition d'un fichier vertical .txt.
    Format attendu: <doc ... edition_id="Edi-XX" ...>
    
    Returns:
        dict: {'id': 'Edi-XX' ou None, 'source': str, 'title': str}
    """
    result = {'id': None, 'source': None, 'title': None}
    
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            # Lire les premières lignes pour trouver la balise <doc>
            for _ in range(50):
                line = f.readline()
                if not line:
                    break
                
                # Chercher edition_id
                match = re.search(r'edition_id="(Edi-\d+)"', line)
                if match:
                    result['id'] = match.group(1)
                
                # Chercher source
                source_match = re.search(r'source="([^"]+)"', line)
                if source_match:
                    result['source'] = source_match.group(1)
                
                # Chercher title
                title_match = re.search(r'title="([^"]+)"', line)
                if title_match:
                    result['title'] = title_match.group(1)
                
                if result['id']:
                    break
    except Exception as e:
        print(f"  Erreur lecture {filepath}: {e}")
    
    return result


def load_libres_de_droits(filepath):
    """
    Charge la liste des identifiants libres de droits depuis un fichier .docx.
    
    Returns:
        set: Ensemble des identifiants (ex: {'Edi-1', 'Edi-15', ...})
    """
    libres = set()
    
    try:
        if filepath.endswith('.docx'):
            doc = Document(filepath)
            text = '\n'.join([p.text for p in doc.paragraphs])
        else:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        
        # Extraire tous les identifiants Edi-XX
        matches = re.findall(r'Edi-(\d+)', text)
        for m in matches:
            libres.add(f"Edi-{m}")
        
        print(f"  → {len(libres)} identifiants libres de droits chargés")
    except Exception as e:
        print(f"  Erreur chargement libres de droits: {e}")
    
    return libres


def scan_sources(fiches_dir, verticaux_dir, textes_dir, verbose=True):
    """
    Scanne les trois sources et retourne un dictionnaire indexé par œuvre (nom du vertical).
    
    Logique :
    - Chaque fichier vertical = une œuvre distincte
    - Les dossiers textes sont nommés d'après les fichiers verticaux
    - Les fiches sont matchées par ID (une fiche peut correspondre à plusieurs œuvres)
    
    Returns:
        dict: {
            'nom_oeuvre_unique': {
                'id': str (Edi-XXX),
                'fiche': chemin ou None,
                'vertical': chemin ou None,
                'textes': chemin du dossier ou None,
                'nom_oeuvre': str,
                'libre_de_droits': bool
            }
        }
    """
    index = {}
    fiches_par_id = {}  # Pour matcher les fiches par ID
    
    debug_stats = {
        'fiches_total': 0,
        'fiches_avec_id': 0,
        'fiches_sans_id': [],
        'verticaux_total': 0,
        'verticaux_avec_id': 0,
        'verticaux_sans_id': [],
        'textes_total': 0,
        'textes_matches': 0,
        'textes_sans_match': []
    }
    
    # 1. Scanner les fiches .docx et les indexer par ID
    if fiches_dir and os.path.isdir(fiches_dir):
        for filepath in glob.glob(os.path.join(fiches_dir, '*.docx')):
            debug_stats['fiches_total'] += 1
            info = extract_info_from_docx(filepath)
            if info['id']:
                debug_stats['fiches_avec_id'] += 1
                if info['id'] not in fiches_par_id:
                    fiches_par_id[info['id']] = {
                        'path': filepath,
                        'titre_oeuvre': info['titre_oeuvre'],
                        'libre_de_droits': info['libre_de_droits']
                    }
            else:
                debug_stats['fiches_sans_id'].append(os.path.basename(filepath))
    
    # 2. Scanner les fichiers verticaux - chaque vertical = une œuvre
    if verticaux_dir and os.path.isdir(verticaux_dir):
        for filepath in glob.glob(os.path.join(verticaux_dir, '*.txt')):
            debug_stats['verticaux_total'] += 1
            info = extract_info_from_vertical(filepath)
            
            # Clé unique = nom du fichier sans extension
            nom_fichier = os.path.splitext(os.path.basename(filepath))[0]
            
            if info['id']:
                debug_stats['verticaux_avec_id'] += 1
                
                # Créer l'entrée pour cette œuvre
                index[nom_fichier] = {
                    'id': info['id'],
                    'fiche': None,
                    'vertical': filepath,
                    'textes': None,
                    'nom_oeuvre': info['source'] or info['title'] or nom_fichier,
                    'libre_de_droits': False
                }
                
                # Associer la fiche si elle existe pour cet ID
                if info['id'] in fiches_par_id:
                    index[nom_fichier]['fiche'] = fiches_par_id[info['id']]['path']
                    if fiches_par_id[info['id']]['libre_de_droits']:
                        index[nom_fichier]['libre_de_droits'] = True
            else:
                debug_stats['verticaux_sans_id'].append(os.path.basename(filepath))
    
    # 3. Scanner les dossiers de textes et les matcher par nom
    if textes_dir and os.path.isdir(textes_dir):
        for item in os.listdir(textes_dir):
            item_path = os.path.join(textes_dir, item)
            if os.path.isdir(item_path):
                debug_stats['textes_total'] += 1
                
                # Le nom du dossier devrait correspondre au nom du fichier vertical
                if item in index:
                    index[item]['textes'] = item_path
                    debug_stats['textes_matches'] += 1
                else:
                    # Essayer sans extension ou avec variations
                    matched = False
                    for key in index.keys():
                        if key == item or key.startswith(item) or item.startswith(key):
                            index[key]['textes'] = item_path
                            debug_stats['textes_matches'] += 1
                            matched = True
                            break
                    if not matched:
                        debug_stats['textes_sans_match'].append(item)
    
    # Afficher les stats de debug
    if verbose:
        print(f"\n  [DEBUG] Scan détaillé :")
        print(f"    Fiches : {debug_stats['fiches_avec_id']}/{debug_stats['fiches_total']} avec ID")
        print(f"    Verticaux : {debug_stats['verticaux_avec_id']}/{debug_stats['verticaux_total']} avec ID")
        print(f"    Textes matchés : {debug_stats['textes_matches']}/{debug_stats['textes_total']}")
        
        if debug_stats['verticaux_sans_id']:
            print(f"\n    ⚠ Verticaux sans ID détecté ({len(debug_stats['verticaux_sans_id'])}) :")
            for f in debug_stats['verticaux_sans_id'][:10]:
                print(f"      - {f}")
            if len(debug_stats['verticaux_sans_id']) > 10:
                print(f"      ... et {len(debug_stats['verticaux_sans_id']) - 10} autres")
        
        if debug_stats['textes_sans_match']:
            print(f"\n    ⚠ Dossiers textes sans match ({len(debug_stats['textes_sans_match'])}) :")
            for f in debug_stats['textes_sans_match'][:10]:
                print(f"      - {f}")
            if len(debug_stats['textes_sans_match']) > 10:
                print(f"      ... et {len(debug_stats['textes_sans_match']) - 10} autres")
    
    return index, debug_stats


def create_export(index, output_dir, libres_de_droits_set=None):
    """
    Crée la structure d'export Nakala.
    Crée un dossier par œuvre (= par fichier vertical).
    Crée un dossier uniquement si vertical ET textes existent.
    
    Returns:
        list: Liste des entrées du log
    """
    if libres_de_droits_set is None:
        libres_de_droits_set = set()
    
    log_entries = []
    stats = {
        'total': 0,
        'libres': 0,
        'non_libres': 0,
        'complets': 0,
        'sans_fiche': 0,
        'fiche_seule': 0,
        'exportes': 0
    }
    
    # Créer les dossiers de base
    libre_dir = os.path.join(output_dir, 'Libre_de_droits')
    non_libre_dir = os.path.join(output_dir, 'Non_libre_de_droits')
    os.makedirs(libre_dir, exist_ok=True)
    os.makedirs(non_libre_dir, exist_ok=True)
    
    for nom_fichier, data in sorted(index.items()):
        stats['total'] += 1
        
        # Déterminer si libre de droits
        is_libre = data['libre_de_droits'] or data['id'] in libres_de_droits_set
        
        if is_libre:
            stats['libres'] += 1
            base_dir = libre_dir
        else:
            stats['non_libres'] += 1
            base_dir = non_libre_dir
        
        # Nom du dossier de sortie
        nom_oeuvre = data['nom_oeuvre'] or nom_fichier
        folder_name = f"{normalize_filename(nom_oeuvre)}_{data['id']}"
        
        # Log entry
        entry = {
            'id': data['id'],
            'nom_fichier': nom_fichier,
            'nom_oeuvre': nom_oeuvre,
            'libre_de_droits': is_libre,
            'export_path': None,
            'exporte': False,
            'has_fiche': data['fiche'] is not None,
            'has_vertical': data['vertical'] is not None,
            'has_textes': data['textes'] is not None,
            'manquants': []
        }
        
        # Vérifier les manquants
        if not data['fiche']:
            entry['manquants'].append('fiche')
        if not data['vertical']:
            entry['manquants'].append('vertical')
        if not data['textes']:
            entry['manquants'].append('textes')
        
        # CRÉER LE DOSSIER SI VERTICAL *ET* TEXTES EXISTENT
        if data['vertical'] and data['textes']:
            export_path = os.path.join(base_dir, folder_name)
            os.makedirs(export_path, exist_ok=True)
            entry['export_path'] = export_path
            entry['exporte'] = True
            stats['exportes'] += 1
            
            # Copier la fiche
            if data['fiche']:
                shutil.copy2(data['fiche'], os.path.join(export_path, 'fiche.docx'))
            
            # Copier le fichier vertical
            shutil.copy2(data['vertical'], os.path.join(export_path, 'vertical.txt'))
            
            # Copier le dossier de textes
            dest_textes = os.path.join(export_path, 'textes')
            shutil.copytree(data['textes'], dest_textes, dirs_exist_ok=True)
            
            # Stats complet/incomplet
            if data['fiche']:
                stats['complets'] += 1
            else:
                stats['sans_fiche'] += 1
        else:
            # Pas de vertical ET textes = pas d'export
            if data['fiche'] or (data['vertical'] and not data['textes']) or (data['textes'] and not data['vertical']):
                stats['fiche_seule'] += 1
        
        log_entries.append(entry)
    
    return log_entries, stats


def write_log(log_entries, stats, debug_stats, output_path):
    """Écrit le fichier de log."""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("RAPPORT D'EXPORT NAKALA (VERSION PAR ŒUVRE)\n")
        f.write(f"Date : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("=" * 80 + "\n\n")
        
        # Statistiques
        f.write("STATISTIQUES\n")
        f.write("-" * 40 + "\n")
        f.write(f"Total d'œuvres trouvées : {stats['total']}\n")
        f.write(f"  - Libres de droits : {stats['libres']}\n")
        f.write(f"  - Non libres de droits : {stats['non_libres']}\n")
        f.write(f"\nExports créés : {stats['exportes']}\n")
        f.write(f"  - Complets (fiche + vertical + textes) : {stats['complets']}\n")
        f.write(f"  - Sans fiche (vertical + textes seulement) : {stats['sans_fiche']}\n")
        f.write(f"\nNon exportés (vertical ou textes manquant) : {stats['fiche_seule']}\n")
        
        # Calculer totaux de fichiers
        total_fiches_copies = sum(1 for e in log_entries if e['exporte'] and e['has_fiche'])
        total_vert_copies = sum(1 for e in log_entries if e['exporte'] and e['has_vertical'])
        total_textes_copies = sum(1 for e in log_entries if e['exporte'] and e['has_textes'])
        f.write(f"\nFichiers copiés :\n")
        f.write(f"  - Fiches : {total_fiches_copies}\n")
        f.write(f"  - Verticaux : {total_vert_copies}\n")
        f.write(f"  - Dossiers textes : {total_textes_copies}\n")
        f.write("\n")
        
        # Debug scan
        f.write("DÉTAIL DU SCAN\n")
        f.write("-" * 40 + "\n")
        f.write(f"Fiches scannées : {debug_stats['fiches_avec_id']}/{debug_stats['fiches_total']} avec ID\n")
        f.write(f"Verticaux scannés : {debug_stats['verticaux_avec_id']}/{debug_stats['verticaux_total']} avec ID\n")
        f.write(f"Dossiers textes matchés : {debug_stats['textes_matches']}/{debug_stats['textes_total']}\n")
        f.write("\n")
        
        # Verticaux sans ID
        if debug_stats['verticaux_sans_id']:
            f.write("=" * 80 + "\n")
            f.write(f"VERTICAUX SANS ID DÉTECTÉ ({len(debug_stats['verticaux_sans_id'])})\n")
            f.write("=" * 80 + "\n")
            for fname in debug_stats['verticaux_sans_id']:
                f.write(f"  - {fname}\n")
            f.write("\n")
        
        # Textes sans match
        if debug_stats['textes_sans_match']:
            f.write("=" * 80 + "\n")
            f.write(f"DOSSIERS TEXTES SANS MATCH ({len(debug_stats['textes_sans_match'])})\n")
            f.write("=" * 80 + "\n")
            for fname in debug_stats['textes_sans_match']:
                f.write(f"  - {fname}\n")
            f.write("\n")
        
        # Œuvres non exportées (vertical ou textes manquant)
        non_exportes = [e for e in log_entries if not e['exporte']]
        if non_exportes:
            f.write("=" * 80 + "\n")
            f.write(f"ŒUVRES NON EXPORTÉES ({len(non_exportes)})\n")
            f.write("=" * 80 + "\n")
            f.write("(Vertical ou textes manquant)\n\n")
            for entry in non_exportes:
                vert_status = "✓" if entry['has_vertical'] else "✗"
                textes_status = "✓" if entry['has_textes'] else "✗"
                fiche_status = "✓" if entry['has_fiche'] else "✗"
                f.write(f"  {entry['id']} - {entry['nom_oeuvre']}\n")
                f.write(f"      Fiche: {fiche_status}  |  Vertical: {vert_status}  |  Textes: {textes_status}\n")
            f.write("\n")
        
        # Exports sans fiche
        sans_fiche = [e for e in log_entries if e['exporte'] and not e['has_fiche']]
        if sans_fiche:
            f.write("=" * 80 + "\n")
            f.write(f"EXPORTS SANS FICHE ({len(sans_fiche)})\n")
            f.write("=" * 80 + "\n")
            for entry in sans_fiche:
                f.write(f"  {entry['id']} - {entry['nom_oeuvre']}\n")
                f.write(f"      Chemin : {entry['export_path']}\n")
            f.write("\n")
        
        # Liste complète des exports par ID
        exportes = [e for e in log_entries if e['exporte']]
        if exportes:
            f.write("=" * 80 + "\n")
            f.write("LISTE DES EXPORTS CRÉÉS\n")
            f.write("=" * 80 + "\n")
            
            # Grouper par ID pour affichage
            par_id = defaultdict(list)
            for entry in exportes:
                par_id[entry['id']].append(entry)
            
            # Trier par numéro d'ID
            def sort_key(id_str):
                try:
                    return int(id_str.split('-')[1])
                except:
                    return 0
            
            # Libres de droits
            libres = [(id_, entries) for id_, entries in par_id.items() if entries[0]['libre_de_droits']]
            if libres:
                f.write("\n[LIBRE DE DROITS]\n")
                for id_, entries in sorted(libres, key=lambda x: sort_key(x[0])):
                    for entry in entries:
                        status = "✓ COMPLET" if entry['has_fiche'] else "⚠ sans fiche"
                        f.write(f"  {entry['id']} - {entry['nom_oeuvre']} : {status}\n")
            
            # Non libres
            non_libres = [(id_, entries) for id_, entries in par_id.items() if not entries[0]['libre_de_droits']]
            if non_libres:
                f.write("\n[NON LIBRE DE DROITS]\n")
                for id_, entries in sorted(non_libres, key=lambda x: sort_key(x[0])):
                    for entry in entries:
                        status = "✓ COMPLET" if entry['has_fiche'] else "⚠ sans fiche"
                        f.write(f"  {entry['id']} - {entry['nom_oeuvre']} : {status}\n")
        
        f.write("\n" + "=" * 80 + "\n")
        f.write("FIN DU RAPPORT\n")


def main(fiches_dir, verticaux_dir, textes_dir, output_dir, libres_path=None):
    """Fonction principale."""
    
    print("=" * 60)
    print("EXPORT NAKALA (VERSION PAR ŒUVRE)")
    print("=" * 60)
    
    # Charger les libres de droits
    libres_set = set()
    if libres_path:
        print(f"\nChargement des libres de droits : {libres_path}")
        libres_set = load_libres_de_droits(libres_path)
    
    # Scanner les sources
    print(f"\nScan des sources...")
    print(f"  Fiches : {fiches_dir}")
    print(f"  Verticaux : {verticaux_dir}")
    print(f"  Textes : {textes_dir}")
    
    index, debug_stats = scan_sources(fiches_dir, verticaux_dir, textes_dir)
    print(f"  → {len(index)} œuvres uniques trouvées")
    
    # Créer l'export
    print(f"\nCréation de l'export dans : {output_dir}")
    os.makedirs(output_dir, exist_ok=True)
    
    log_entries, stats = create_export(index, output_dir, libres_set)
    
    # Écrire le log
    log_path = os.path.join(output_dir, 'log_export.txt')
    write_log(log_entries, stats, debug_stats, log_path)
    
    # Résumé
    print("\n" + "=" * 60)
    print("RÉSUMÉ")
    print("=" * 60)
    print(f"Total trouvé : {stats['total']} œuvres")
    print(f"  - Libres de droits : {stats['libres']}")
    print(f"  - Non libres de droits : {stats['non_libres']}")
    print(f"\nExports créés : {stats['exportes']}")
    print(f"  - Complets (fiche + vertical + textes) : {stats['complets']}")
    print(f"  - Sans fiche (vertical + textes) : {stats['sans_fiche']}")
    print(f"\nNon exportés (vertical ou textes manquant) : {stats['fiche_seule']}")
    print(f"\nLog : {log_path}")
    
    return log_entries, stats


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Crée l'export Nakala (version par œuvre) en associant fiches, verticaux et textes."
    )
    parser.add_argument('--fiches', '-f', required=True, help='Dossier des fiches .docx')
    parser.add_argument('--verticaux', '-v', required=True, help='Dossier des fichiers verticaux .txt')
    parser.add_argument('--textes', '-t', required=True, help='Dossier des sous-dossiers de textes')
    parser.add_argument('--output', '-o', default='Export_Nakala', help='Dossier de sortie')
    parser.add_argument('--libres', '-l', default=None, help='Fichier listant les libres de droits')
    
    args = parser.parse_args()
    
    main(args.fiches, args.verticaux, args.textes, args.output, args.libres)
