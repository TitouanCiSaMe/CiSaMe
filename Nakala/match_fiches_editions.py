#!/usr/bin/env python3
"""
Script pour faire correspondre les fiches d'œuvres (.docx) avec la base de données
des éditions (CSV) et ajouter l'identifiant de l'édition à la fin de chaque fiche.

Stratégie de matching (par ordre de priorité) :
1. Correspondance sur le titre de l'édition
2. Correspondance sur le titre de l'œuvre + auteur
3. Correspondance fuzzy si pas de match exact

Auteur: Script pour le projet CiSaMe
"""

import csv
import re
import sys
import os
import logging
from pathlib import Path
from difflib import SequenceMatcher
from docx import Document
from nakala_utils import normalize_text

logger = logging.getLogger(__name__)


def load_libres_de_droits(filepath):
    """
    Charge la liste des identifiants ET titres libres de droits depuis un fichier .docx ou .txt.
    Retourne un dict avec:
        - 'identifiants': set d'identifiants (ex: {'Edi-1', 'Edi-15', ...})
        - 'titres': set de titres normalisés (pour les œuvres sans ID)
    """
    result = {
        'identifiants': set(),
        'titres': set()
    }
    
    if filepath.endswith('.docx'):
        doc = Document(filepath)
        text = '\n'.join([p.text for p in doc.paragraphs])
    else:
        with open(filepath, 'r', encoding='utf-8') as f:
            text = f.read()
    
    # Extraire tous les identifiants Edi-XX
    matches = re.findall(r'Edi-(\d+)', text)
    for m in matches:
        result['identifiants'].add(f"Edi-{m}")
    
    # Extraire les titres des lignes avec "?" ou sans Edi-
    for line in text.split('\n'):
        line = line.strip()
        if not line or line.startswith('==>'):
            continue
        
        # Lignes avec "/ ?" = œuvre sans ID connu
        if '/ ?' in line or ('/' in line and 'Edi-' not in line and '?' in line):
            # Format typique: "Auteur, Titre / ?" ou "Titre / ?"
            if '/' in line:
                before_slash = line.split('/')[0].strip()
                # Si virgule, le titre est après la virgule
                if ',' in before_slash:
                    titre = before_slash.split(',', 1)[1].strip()
                else:
                    titre = before_slash
                # Nettoyer le titre
                titre = titre.replace('\xa0', ' ').strip()
                if titre:
                    result['titres'].add(normalize_text(titre))
    
    return result


def extract_info_from_docx(filepath):
    """
    Extrait les informations clés d'une fiche docx.
    Retourne un dictionnaire avec les champs trouvés.

    Raises:
        IOError: Si le fichier ne peut pas être lu.
    """
    try:
        doc = Document(filepath)
    except Exception as e:
        logger.warning("Impossible de lire %s : %s", filepath, e)
        raise IOError(f"Impossible de lire {filepath}") from e
    text = '\n'.join([p.text for p in doc.paragraphs])
    
    info = {
        'titre_oeuvre': None,
        'auteur': None,
        'titre_edition': None,
        'editeur': None,
        'date_edition': None,
        'lieu_edition': None,
        'raw_text': text
    }
    
    # Patterns pour extraire les champs
    # Titre de l'œuvre (premier "Titre :" avant "Édition :")
    edition_pos = text.find('Édition')
    if edition_pos == -1:
        edition_pos = len(text)
    
    titre_match = re.search(r'Titre\s*:\s*(.+?)(?:\n|$)', text[:edition_pos])
    if titre_match:
        info['titre_oeuvre'] = titre_match.group(1).strip().strip('*').strip()
    
    # Auteur(s)
    auteur_match = re.search(r'Auteur\(s\)\s*:\s*(.+?)(?:\n|$)', text)
    if auteur_match:
        # Prendre le premier nom (avant les parenthèses ou le premier ";")
        auteur_full = auteur_match.group(1).strip()
        # Extraire le nom principal (avant les variantes entre parenthèses)
        auteur_principal = re.split(r'\s*\(|\s*;', auteur_full)[0].strip()
        info['auteur'] = auteur_principal
    
    # Informations de l'édition (après "Édition :")
    edition_section = text[edition_pos:] if edition_pos < len(text) else ""
    
    # Titre de l'édition
    titre_ed_match = re.search(r'Titre\s*:\s*(.+?)(?:\n|$)', edition_section)
    if titre_ed_match:
        info['titre_edition'] = titre_ed_match.group(1).strip().strip('*').strip()
    
    # Éditeur
    editeur_match = re.search(r'Éditeur\s*:\s*(.+?)(?:\n|$)', edition_section)
    if editeur_match:
        info['editeur'] = editeur_match.group(1).strip()
    
    # Date d'édition
    date_match = re.search(r"Date d'édition\s*:\s*(\d{4})", edition_section)
    if date_match:
        info['date_edition'] = date_match.group(1)
    
    # Lieu d'édition
    lieu_match = re.search(r"Lieu d'édition\s*:\s*(.+?)(?:\n|$)", edition_section)
    if lieu_match:
        info['lieu_edition'] = lieu_match.group(1).strip()
    
    return info


def load_csv_editions(csv_path):
    """
    Charge le fichier CSV des éditions.
    Retourne une liste de dictionnaires avec les informations clés.
    
    Note: Le CSV a des colonnes avec le même nom, donc on utilise les index.
    Index des colonnes importantes :
    0: Edition H-ID
    1: Edition Record Title
    2: Identifiant interne (de l'édition, ex: Edi-39)
    3: Title
    5: Éditeur
    8: Date
    11: Oeuvre Record Title
    12: Identifiant interne (de l'œuvre, ex: Oeuv-39)
    15: Person Record Title (auteur)
    28: Lieu d'édition
    """
    editions = []
    
    with open(csv_path, 'r', encoding='utf-8-sig') as f:
        # Le fichier est séparé par des tabulations
        reader = csv.reader(f, delimiter='\t')
        
        # Sauter l'en-tête
        next(reader)
        
        for row in reader:
            if len(row) < 20:  # Ignorer les lignes trop courtes
                continue
            
            edition = {
                'edition_hid': row[0] if len(row) > 0 else '',
                'edition_title': row[1] if len(row) > 1 else '',
                'identifiant_edition': row[2] if len(row) > 2 else '',  # Ex: Edi-39
                'title': row[3] if len(row) > 3 else '',
                'editeur': row[5] if len(row) > 5 else '',
                'date': row[8] if len(row) > 8 else '',
                'oeuvre_title': row[11] if len(row) > 11 else '',
                'identifiant_oeuvre': row[12] if len(row) > 12 else '',  # Ex: Oeuv-39
                'auteur': row[15] if len(row) > 15 else '',
                'lieu_edition': row[28] if len(row) > 28 else '',
                # Normaliser pour la recherche
                'norm_edition_title': normalize_text(row[1] if len(row) > 1 else ''),
                'norm_oeuvre_title': normalize_text(row[11] if len(row) > 11 else ''),
                'norm_auteur': normalize_text(row[15] if len(row) > 15 else ''),
            }
            editions.append(edition)
    
    return editions


def similarity_score(s1, s2):
    """Calcule un score de similarité entre deux chaînes (0-1)."""
    if not s1 or not s2:
        return 0.0
    return SequenceMatcher(None, normalize_text(s1), normalize_text(s2)).ratio()


def find_best_match(fiche_info, editions, threshold=0.7):
    """
    Trouve la meilleure correspondance dans la base CSV pour une fiche.
    
    Stratégie :
    1. Match exact sur titre d'édition + date
    2. Match exact sur titre d'œuvre + auteur
    3. Match fuzzy avec score le plus élevé
    
    Retourne (edition_match, score, method) ou (None, 0, None) si pas de match.
    """
    best_match = None
    best_score = 0
    match_method = None
    
    norm_titre_oeuvre = normalize_text(fiche_info['titre_oeuvre']) if fiche_info['titre_oeuvre'] else ""
    norm_titre_edition = normalize_text(fiche_info['titre_edition']) if fiche_info['titre_edition'] else ""
    norm_auteur = normalize_text(fiche_info['auteur']) if fiche_info['auteur'] else ""
    
    for edition in editions:
        score = 0
        method = ""
        
        # 1. Match sur titre d'édition
        if norm_titre_edition and edition['norm_edition_title']:
            title_sim = similarity_score(fiche_info['titre_edition'], edition['edition_title'])
            if title_sim > 0.8:
                score = title_sim
                method = "titre_edition"
                # Bonus si la date correspond aussi
                if fiche_info['date_edition'] and fiche_info['date_edition'] in str(edition['date']):
                    score += 0.2
                    method += "+date"
        
        # 2. Match sur titre d'œuvre + auteur
        if norm_titre_oeuvre:
            oeuvre_sim = similarity_score(fiche_info['titre_oeuvre'], edition['oeuvre_title'])
            if oeuvre_sim > 0.8:
                current_score = oeuvre_sim
                current_method = "titre_oeuvre"
                
                # Bonus pour l'auteur
                if norm_auteur and edition['norm_auteur']:
                    auteur_sim = similarity_score(fiche_info['auteur'], edition['auteur'])
                    if auteur_sim > 0.6:
                        current_score += 0.3
                        current_method += "+auteur"
                
                # Bonus pour la date
                if fiche_info['date_edition'] and fiche_info['date_edition'] in str(edition['date']):
                    current_score += 0.2
                    current_method += "+date"
                
                if current_score > score:
                    score = current_score
                    method = current_method
        
        # Mise à jour du meilleur match
        if score > best_score:
            best_score = score
            best_match = edition
            match_method = method
    
    # Plafonner le score a 1.0
    best_score = min(best_score, 1.0)

    # Verifier si le score depasse le seuil
    if best_score >= threshold:
        return best_match, best_score, match_method

    return None, 0, None


def add_identifier_to_docx(docx_path, identifier, output_path=None, libre_de_droits=False):
    """
    Ajoute l'identifiant de l'édition à la fin du document docx.
    
    Args:
        docx_path: Chemin du fichier docx source
        identifier: Identifiant à ajouter (ex: "Edi-39")
        output_path: Chemin de sortie (si None, modifie le fichier en place)
        libre_de_droits: Si True, ajoute aussi la mention "Libre de droits : Oui"
    """
    doc = Document(docx_path)
    
    # Ajouter une ligne vide puis l'identifiant
    doc.add_paragraph()
    id_para = doc.add_paragraph()
    id_run = id_para.add_run(f"Identifiant édition : {identifier}")
    id_run.bold = True
    
    # Ajouter la mention libre de droits si applicable
    if libre_de_droits:
        libre_para = doc.add_paragraph()
        libre_run = libre_para.add_run("Libre de droits : Oui")
        libre_run.bold = True
    
    save_path = output_path or docx_path
    doc.save(save_path)
    return save_path


def process_fiche(docx_path, csv_editions, output_dir=None, libres_de_droits=None, verbose=True):
    """
    Traite une fiche : extraction, matching, ajout de l'identifiant.
    
    Args:
        docx_path: Chemin du fichier fiche docx
        csv_editions: Liste des éditions chargées depuis le CSV
        output_dir: Répertoire de sortie (si None, modifie en place)
        libres_de_droits: Dict avec 'identifiants' et 'titres' libres de droits (optionnel)
        verbose: Afficher les détails
    
    Returns:
        dict avec les résultats du traitement
    """
    if libres_de_droits is None:
        libres_de_droits = {'identifiants': set(), 'titres': set()}
    
    result = {
        'fichier': os.path.basename(docx_path),
        'succes': False,
        'match': None,
        'score': 0,
        'methode': None,
        'identifiant': None,
        'libre_de_droits': False,
        'output_path': None,
        'erreur': None
    }
    
    try:
        # 1. Extraire les informations de la fiche
        fiche_info = extract_info_from_docx(docx_path)
        
        if verbose:
            print(f"\n{'='*60}")
            print(f"Traitement de : {os.path.basename(docx_path)}")
            print(f"  Titre œuvre : {fiche_info['titre_oeuvre']}")
            print(f"  Auteur : {fiche_info['auteur']}")
            print(f"  Titre édition : {fiche_info['titre_edition']}")
            print(f"  Date édition : {fiche_info['date_edition']}")
        
        # 2. Trouver la correspondance
        match, score, method = find_best_match(fiche_info, csv_editions)
        
        if match:
            result['succes'] = True
            result['match'] = match
            result['score'] = score
            result['methode'] = method
            result['identifiant'] = match['identifiant_edition']
            
            # Vérifier si libre de droits
            # 1. Par identifiant
            is_libre = match['identifiant_edition'] in libres_de_droits.get('identifiants', set())
            
            # 2. Par titre de l'œuvre (si pas trouvé par ID)
            if not is_libre and fiche_info['titre_oeuvre']:
                norm_titre = normalize_text(fiche_info['titre_oeuvre'])
                for titre_libre in libres_de_droits.get('titres', set()):
                    # Match si le titre contient ou est contenu
                    if titre_libre in norm_titre or norm_titre in titre_libre:
                        is_libre = True
                        break
                    # Ou similarité élevée
                    if similarity_score(norm_titre, titre_libre) > 0.8:
                        is_libre = True
                        break
            
            # 3. Par titre de l'œuvre dans le CSV
            if not is_libre and match['oeuvre_title']:
                norm_titre_csv = normalize_text(match['oeuvre_title'])
                for titre_libre in libres_de_droits.get('titres', set()):
                    if titre_libre in norm_titre_csv or norm_titre_csv in titre_libre:
                        is_libre = True
                        break
                    if similarity_score(norm_titre_csv, titre_libre) > 0.8:
                        is_libre = True
                        break
            
            result['libre_de_droits'] = is_libre
            
            if verbose:
                print(f"\n  ✓ MATCH TROUVÉ (score: {score:.2f}, méthode: {method})")
                print(f"    → Édition : {match['edition_title'][:60]}...")
                print(f"    → Œuvre : {match['oeuvre_title']}")
                print(f"    → Identifiant : {match['identifiant_edition']}")
                if is_libre:
                    print(f"    → Libre de droits : Oui ✓")
            
            # 3. Ajouter l'identifiant au document
            if output_dir:
                output_path = os.path.join(output_dir, os.path.basename(docx_path))
            else:
                # Créer un nouveau fichier avec suffixe _id
                base, ext = os.path.splitext(docx_path)
                output_path = f"{base}_avec_id{ext}"
            
            result['output_path'] = add_identifier_to_docx(
                docx_path, 
                match['identifiant_edition'], 
                output_path,
                libre_de_droits=is_libre
            )
            
            if verbose:
                print(f"    → Fichier sauvegardé : {os.path.basename(result['output_path'])}")
        else:
            if verbose:
                print(f"\n  ✗ Aucune correspondance trouvée (seuil: 0.7)")
    
    except Exception as e:
        result['erreur'] = str(e)
        if verbose:
            print(f"\n  ✗ ERREUR : {e}")
    
    return result


def main(docx_files, csv_path, output_dir=None, libres_de_droits_path=None):
    """
    Fonction principale pour traiter plusieurs fiches.
    
    Args:
        docx_files: Liste des chemins des fiches docx
        csv_path: Chemin du fichier CSV des éditions
        output_dir: Répertoire de sortie (optionnel)
        libres_de_droits_path: Chemin vers le fichier listant les éditions libres de droits
    """
    print("Chargement de la base des éditions...")
    editions = load_csv_editions(csv_path)
    print(f"  → {len(editions)} éditions chargées")
    
    # Charger la liste des libres de droits si fournie
    libres_de_droits = {'identifiants': set(), 'titres': set()}
    if libres_de_droits_path:
        print(f"Chargement de la liste des libres de droits...")
        libres_de_droits = load_libres_de_droits(libres_de_droits_path)
        print(f"  → {len(libres_de_droits['identifiants'])} identifiants (Edi-XX)")
        print(f"  → {len(libres_de_droits['titres'])} titres sans identifiant")
    
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
    
    results = []
    for docx_path in docx_files:
        result = process_fiche(docx_path, editions, output_dir, libres_de_droits)
        results.append(result)
    
    # Résumé
    print("\n" + "="*60)
    print("RÉSUMÉ")
    print("="*60)
    success = sum(1 for r in results if r['succes'])
    libres = sum(1 for r in results if r.get('libre_de_droits'))
    print(f"Fiches traitées : {len(results)}")
    print(f"Correspondances trouvées : {success}/{len(results)}")
    if libres_de_droits:
        print(f"Dont libres de droits : {libres}")
    
    if success < len(results):
        print("\nFiches sans correspondance :")
        for r in results:
            if not r['succes']:
                print(f"  - {r['fichier']}: {r.get('erreur', 'Pas de match')}")
    
    return results


# Exemple d'utilisation
if __name__ == "__main__":
    import argparse
    import glob
    
    parser = argparse.ArgumentParser(
        description="Fait correspondre les fiches d'œuvres (.docx) avec la base CSV et ajoute l'identifiant de l'édition."
    )
    parser.add_argument('--csv', required=True, help='Chemin vers le fichier CSV des éditions')
    parser.add_argument('--output', '-o', default=None, help='Répertoire de sortie (optionnel)')
    parser.add_argument('--dossier', '-d', default=None, help='Dossier contenant les fiches .docx')
    parser.add_argument('--libres', '-l', default=None, help='Fichier .docx listant les éditions libres de droits')
    parser.add_argument('fiches', nargs='*', help='Fichiers .docx des fiches à traiter')
    
    args = parser.parse_args()
    
    # Collecter les fichiers
    docx_files = []
    
    # Option 1: dossier spécifié
    if args.dossier:
        dossier = args.dossier.rstrip('/')
        docx_files.extend(glob.glob(f"{dossier}/*.docx"))
        docx_files.extend(glob.glob(f"{dossier}/*.DOCX"))
    
    # Option 2: fichiers listés explicitement
    if args.fiches:
        docx_files.extend(args.fiches)
    
    if not docx_files:
        print("Erreur: Aucun fichier .docx trouvé.")
        print("Utilisation:")
        print("  python match_fiches_editions.py --csv export.csv --dossier ./fiches/")
        print("  python match_fiches_editions.py --csv export.csv --libres libres.docx --dossier ./fiches/")
        sys.exit(1)
    
    # Supprimer les doublons et trier
    docx_files = sorted(set(docx_files))
    print(f"Fichiers à traiter : {len(docx_files)}")
    
    results = main(docx_files, args.csv, args.output, args.libres)
    
    # Code de sortie
    sys.exit(0 if all(r['succes'] for r in results) else 1)
