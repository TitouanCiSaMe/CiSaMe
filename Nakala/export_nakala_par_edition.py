#!/usr/bin/env python3
"""
Export Nakala - Version par ÉDITION (regroupée)

Logique : Un dossier = un ID d'édition (Edi-XX)
Si plusieurs œuvres partagent le même ID, elles sont regroupées dans le même dossier.

Résultat attendu : ~87 dossiers (pour 122 fichiers si certains partagent des IDs)

Structure de sortie :
    Export_Nakala/
    ├── Libre_de_droits/
    │   └── Nom_oeuvre_Edi-XX/
    │       ├── fiche.docx
    │       ├── vertical_1.txt
    │       ├── vertical_2.txt      (si plusieurs œuvres)
    │       ├── textes_1/
    │       └── textes_2/
    └── Non_libre_de_droits/
        └── ...

Auteur: Script pour le projet CiSaMe
"""

import os
import re
import glob
import logging
import shutil
from datetime import datetime
from docx import Document
from nakala_utils import (
    normalize_filename, extract_info_from_docx, extract_info_from_vertical,
    load_libres_de_droits, edi_sort_key,
)

logger = logging.getLogger(__name__)


def normalize_filename(name, max_length=80):
    """Normalise un nom pour l'utiliser comme nom de dossier."""
    if not name:
        return "Inconnu"
    # Remplacer les caractères problématiques
    name = re.sub(r'[<>:"/\\|?*]', '_', name)
    name = re.sub(r'\s+', '_', name)
    name = re.sub(r'_+', '_', name)
    name = name.strip('_')
    # Limiter la longueur
    if len(name) > max_length:
        name = name[:max_length].rstrip('_')
    return name


def extract_info_from_docx(filepath):
    """
    Extrait l'identifiant édition et le statut libre de droits d'une fiche .docx.
    
    Returns:
        dict: {'id': 'Edi-XX' ou None, 'titre_oeuvre': str, 'libre_de_droits': bool}
    """
    result = {'id': None, 'titre_oeuvre': None, 'libre_de_droits': False}
    
    try:
        doc = Document(filepath)
        text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Chercher l'identifiant édition
        match = re.search(r'Identifiant\s+[ée]dition\s*:\s*(Edi-\d+)', text, re.IGNORECASE)
        if match:
            result['id'] = match.group(1)
        
        # Chercher le titre de l'œuvre
        titre_match = re.search(r'Titre\s*:\s*\*?([^*\n]+)\*?', text)
        if titre_match:
            result['titre_oeuvre'] = titre_match.group(1).strip()
        
        # Chercher si libre de droits
        if re.search(r'Libre\s+de\s+droits\s*:\s*Oui', text, re.IGNORECASE):
            result['libre_de_droits'] = True
    except (IOError, OSError, ValueError) as e:
        logger.warning("Erreur lecture %s : %s", filepath, e)
    
    return result


def extract_info_from_vertical(filepath):
    """
    Extrait l'identifiant édition d'un fichier vertical .txt.
    Format attendu: <doc ... edition_id="Edi-XX" ...>
    
    Returns:
        dict: {'id': 'Edi-XX' ou None, 'source': str, 'title': str}
    """
    result = {'id': None, 'source': None, 'title': None}
    
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            # Lire les premières lignes pour trouver la balise <doc>
            for _ in range(50):
                line = f.readline()
                if not line:
                    break
                
                # Chercher edition_id
                match = re.search(r'edition_id="(Edi-\d+)"', line)
                if match:
                    result['id'] = match.group(1)
                
                # Chercher source
                source_match = re.search(r'source="([^"]+)"', line)
                if source_match:
                    result['source'] = source_match.group(1)
                
                # Chercher title
                title_match = re.search(r'title="([^"]+)"', line)
                if title_match:
                    result['title'] = title_match.group(1)
                
                if result['id']:
                    break
    except (IOError, OSError, ValueError) as e:
        logger.warning("Erreur lecture %s : %s", filepath, e)
    
    return result


def extract_info_from_textes(dirpath):
    """
    Extrait l'identifiant édition d'un dossier de textes.
    Cherche dans les fichiers du dossier.
    
    Returns:
        dict: {'id': 'Edi-XX' ou None, 'provenance': str, 'oeuvre': str}
    """
    result = {'id': None, 'provenance': None, 'oeuvre': None}
    
    try:
        # Chercher d'abord un fichier fusion ou le premier fichier txt
        files = []
        for f in os.listdir(dirpath):
            if f.endswith('.txt'):
                if 'fusion' in f.lower() or 'complet' in f.lower():
                    files.insert(0, f)  # Priorité aux fichiers de fusion
                else:
                    files.append(f)
        
        for fname in files[:5]:  # Vérifier les 5 premiers fichiers max
            fpath = os.path.join(dirpath, fname)
            with open(fpath, 'r', encoding='utf-8') as f:
                content = f.read(2000)  # Lire le début
                
                # Chercher Edition ID
                match = re.search(r'Edition\s+ID\s*:\s*(Edi-\d+)', content)
                if match:
                    result['id'] = match.group(1)
                
                # Chercher Provenance
                prov_match = re.search(r'Provenance\s*:\s*([^\n]+)', content)
                if prov_match:
                    result['provenance'] = prov_match.group(1).strip()
                
                # Chercher Œuvre
                oeuvre_match = re.search(r'Œuvre\s*:\s*([^\n]+)', content)
                if oeuvre_match:
                    result['oeuvre'] = oeuvre_match.group(1).strip()
                
                if result['id']:
                    break
    except (IOError, OSError, ValueError) as e:
        logger.warning("Erreur lecture %s : %s", dirpath, e)
    
    return result


def load_libres_de_droits(filepath):
    """
    Charge la liste des identifiants libres de droits depuis un fichier .docx.
    
    Returns:
        set: Ensemble des identifiants (ex: {'Edi-1', 'Edi-15', ...})
    """
    libres = set()
    
    try:
        if filepath.endswith('.docx'):
            doc = Document(filepath)
            text = '\n'.join([p.text for p in doc.paragraphs])
        else:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        
        # Extraire tous les identifiants Edi-XX
        matches = re.findall(r'Edi-(\d+)', text)
        for m in matches:
            libres.add(f"Edi-{m}")
        
        print(f"  → {len(libres)} identifiants libres de droits chargés")
    except (IOError, OSError, ValueError) as e:
        logger.warning("Erreur chargement libres de droits : %s", e)
    
    return libres


def scan_sources(fiches_dir, verticaux_dir, textes_dir, verbose=True):
    """
    Scanne les trois sources et retourne un dictionnaire indexé par Edi-XXX.
    
    Returns:
        dict: {
            'Edi-XX': {
                'fiches': [liste de chemins],
                'verticaux': [liste de chemins],
                'textes': [liste de chemins de dossiers],
                'nom_oeuvre': str,
                'libre_de_droits': bool
            }
        }
    """
    index = {}
    debug_stats = {
        'fiches_total': 0,
        'fiches_avec_id': 0,
        'fiches_sans_id': [],
        'verticaux_total': 0,
        'verticaux_avec_id': 0,
        'verticaux_sans_id': [],
        'textes_total': 0,
        'textes_avec_id': 0,
        'textes_sans_id': []
    }
    
    # 1. Scanner les fiches .docx
    if fiches_dir and os.path.isdir(fiches_dir):
        for filepath in glob.glob(os.path.join(fiches_dir, '*.docx')):
            debug_stats['fiches_total'] += 1
            info = extract_info_from_docx(filepath)
            if info['id']:
                debug_stats['fiches_avec_id'] += 1
                if info['id'] not in index:
                    index[info['id']] = {
                        'fiches': [], 'verticaux': [], 'textes': [],
                        'nom_oeuvre': None, 'libre_de_droits': False
                    }
                index[info['id']]['fiches'].append(filepath)
                if info['titre_oeuvre']:
                    index[info['id']]['nom_oeuvre'] = info['titre_oeuvre']
                if info['libre_de_droits']:
                    index[info['id']]['libre_de_droits'] = True
            else:
                debug_stats['fiches_sans_id'].append(os.path.basename(filepath))
    
    # 2. Scanner les fichiers verticaux .txt
    if verticaux_dir and os.path.isdir(verticaux_dir):
        for filepath in glob.glob(os.path.join(verticaux_dir, '*.txt')):
            debug_stats['verticaux_total'] += 1
            info = extract_info_from_vertical(filepath)
            if info['id']:
                debug_stats['verticaux_avec_id'] += 1
                if info['id'] not in index:
                    index[info['id']] = {
                        'fiches': [], 'verticaux': [], 'textes': [],
                        'nom_oeuvre': None, 'libre_de_droits': False
                    }
                index[info['id']]['verticaux'].append(filepath)
                # Utiliser source ou title comme nom d'œuvre si pas déjà défini
                if not index[info['id']]['nom_oeuvre']:
                    index[info['id']]['nom_oeuvre'] = info['source'] or info['title']
            else:
                debug_stats['verticaux_sans_id'].append(os.path.basename(filepath))
    
    # 3. Scanner les dossiers de textes
    if textes_dir and os.path.isdir(textes_dir):
        for item in os.listdir(textes_dir):
            item_path = os.path.join(textes_dir, item)
            if os.path.isdir(item_path):
                debug_stats['textes_total'] += 1
                info = extract_info_from_textes(item_path)
                if info['id']:
                    debug_stats['textes_avec_id'] += 1
                    if info['id'] not in index:
                        index[info['id']] = {
                            'fiches': [], 'verticaux': [], 'textes': [],
                            'nom_oeuvre': None, 'libre_de_droits': False
                        }
                    index[info['id']]['textes'].append(item_path)
                    # Utiliser provenance ou œuvre comme nom si pas défini
                    if not index[info['id']]['nom_oeuvre']:
                        index[info['id']]['nom_oeuvre'] = info['provenance'] or info['oeuvre']
                else:
                    debug_stats['textes_sans_id'].append(item)
    
    # Afficher les stats de debug
    if verbose:
        print(f"\n  [DEBUG] Scan détaillé :")
        print(f"    Fiches : {debug_stats['fiches_avec_id']}/{debug_stats['fiches_total']} avec ID")
        print(f"    Verticaux : {debug_stats['verticaux_avec_id']}/{debug_stats['verticaux_total']} avec ID")
        print(f"    Textes : {debug_stats['textes_avec_id']}/{debug_stats['textes_total']} avec ID")
        
        if debug_stats['verticaux_sans_id']:
            print(f"\n    ⚠ Verticaux sans ID détecté ({len(debug_stats['verticaux_sans_id'])}) :")
            for f in debug_stats['verticaux_sans_id'][:10]:
                print(f"      - {f}")
            if len(debug_stats['verticaux_sans_id']) > 10:
                print(f"      ... et {len(debug_stats['verticaux_sans_id']) - 10} autres")
        
        if debug_stats['textes_sans_id']:
            print(f"\n    ⚠ Dossiers textes sans ID détecté ({len(debug_stats['textes_sans_id'])}) :")
            for f in debug_stats['textes_sans_id'][:10]:
                print(f"      - {f}")
            if len(debug_stats['textes_sans_id']) > 10:
                print(f"      ... et {len(debug_stats['textes_sans_id']) - 10} autres")
    
    return index, debug_stats


def create_export(index, output_dir, libres_de_droits_set=None):
    """
    Crée la structure d'export Nakala.
    Crée un dossier uniquement si vertical ET textes existent.
    
    Returns:
        list: Liste des entrées du log
    """
    if libres_de_droits_set is None:
        libres_de_droits_set = set()
    
    log_entries = []
    stats = {
        'total': 0,
        'libres': 0,
        'non_libres': 0,
        'complets': 0,
        'sans_fiche': 0,
        'fiche_seule': 0,
        'exportes': 0
    }
    
    # Créer les dossiers de base
    libre_dir = os.path.join(output_dir, 'Libre_de_droits')
    non_libre_dir = os.path.join(output_dir, 'Non_libre_de_droits')
    os.makedirs(libre_dir, exist_ok=True)
    os.makedirs(non_libre_dir, exist_ok=True)
    
    for edi_id, data in sorted(index.items(), key=lambda x: edi_sort_key(x[0])):
        stats['total'] += 1
        
        # Déterminer si libre de droits
        is_libre = data['libre_de_droits'] or edi_id in libres_de_droits_set
        
        if is_libre:
            stats['libres'] += 1
            base_dir = libre_dir
        else:
            stats['non_libres'] += 1
            base_dir = non_libre_dir
        
        # Nom du dossier de sortie
        nom_oeuvre = data['nom_oeuvre'] or 'Oeuvre_inconnue'
        folder_name = f"{normalize_filename(nom_oeuvre)}_{edi_id}"
        
        # Log entry
        entry = {
            'id': edi_id,
            'nom_oeuvre': nom_oeuvre,
            'libre_de_droits': is_libre,
            'export_path': None,
            'exporte': False,
            'fiches': len(data['fiches']),
            'verticaux': len(data['verticaux']),
            'textes': len(data['textes']),
            'manquants': []
        }
        
        # Vérifier les manquants
        if not data['fiches']:
            entry['manquants'].append('fiche')
        if not data['verticaux']:
            entry['manquants'].append('vertical')
        if not data['textes']:
            entry['manquants'].append('textes')
        
        # CRÉER LE DOSSIER SI VERTICAL *ET* TEXTES EXISTENT
        if data['verticaux'] and data['textes']:
            export_path = os.path.join(base_dir, folder_name)
            os.makedirs(export_path, exist_ok=True)
            entry['export_path'] = export_path
            entry['exporte'] = True
            stats['exportes'] += 1
            
            # Copier les fiches
            for i, fiche_path in enumerate(data['fiches']):
                suffix = f"_{i+1}" if len(data['fiches']) > 1 else ""
                dest_name = f"fiche{suffix}.docx"
                shutil.copy2(fiche_path, os.path.join(export_path, dest_name))
            
            # Copier les fichiers verticaux
            for i, vert_path in enumerate(data['verticaux']):
                suffix = f"_{i+1}" if len(data['verticaux']) > 1 else ""
                dest_name = f"vertical{suffix}.txt"
                shutil.copy2(vert_path, os.path.join(export_path, dest_name))
            
            # Copier les dossiers de textes
            for i, textes_path in enumerate(data['textes']):
                suffix = f"_{i+1}" if len(data['textes']) > 1 else ""
                dest_name = f"textes{suffix}"
                dest_path = os.path.join(export_path, dest_name)
                shutil.copytree(textes_path, dest_path, dirs_exist_ok=True)
            
            # Stats complet/incomplet
            if data['fiches']:
                stats['complets'] += 1
            else:
                stats['sans_fiche'] += 1
        else:
            # Pas de vertical ET textes = pas d'export
            if data['fiches']:
                stats['fiche_seule'] += 1
        
        log_entries.append(entry)
    
    return log_entries, stats


def write_log(log_entries, stats, debug_stats, output_path):
    """Écrit le fichier de log."""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("RAPPORT D'EXPORT NAKALA (VERSION PAR ÉDITION)\n")
        f.write(f"Date : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("=" * 80 + "\n\n")
        
        # Statistiques
        f.write("STATISTIQUES\n")
        f.write("-" * 40 + "\n")
        f.write(f"Total d'éditions trouvées : {stats['total']}\n")
        f.write(f"  - Libres de droits : {stats['libres']}\n")
        f.write(f"  - Non libres de droits : {stats['non_libres']}\n")
        f.write(f"\nExports créés : {stats['exportes']}\n")
        f.write(f"  - Complets (fiche + vertical + textes) : {stats['complets']}\n")
        f.write(f"  - Sans fiche (vertical + textes seulement) : {stats['sans_fiche']}\n")
        f.write(f"\nNon exportés (fiche seule, sans vertical ou textes) : {stats['fiche_seule']}\n")
        
        # Calculer totaux de fichiers
        total_fiches_copies = sum(e['fiches'] for e in log_entries if e['exporte'])
        total_vert_copies = sum(e['verticaux'] for e in log_entries if e['exporte'])
        total_textes_copies = sum(e['textes'] for e in log_entries if e['exporte'])
        f.write(f"\nFichiers copiés :\n")
        f.write(f"  - Fiches : {total_fiches_copies}\n")
        f.write(f"  - Verticaux : {total_vert_copies}\n")
        f.write(f"  - Dossiers textes : {total_textes_copies}\n")
        f.write("\n")
        
        # Debug scan
        f.write("DÉTAIL DU SCAN\n")
        f.write("-" * 40 + "\n")
        f.write(f"Fiches scannées : {debug_stats['fiches_avec_id']}/{debug_stats['fiches_total']} avec ID\n")
        f.write(f"Verticaux scannés : {debug_stats['verticaux_avec_id']}/{debug_stats['verticaux_total']} avec ID\n")
        f.write(f"Dossiers textes scannés : {debug_stats['textes_avec_id']}/{debug_stats['textes_total']} avec ID\n")
        f.write("\n")
        
        # Éditions avec plusieurs verticaux
        multi_verticaux = [e for e in log_entries if e['verticaux'] > 1]
        if multi_verticaux:
            f.write("=" * 80 + "\n")
            f.write(f"ÉDITIONS AVEC PLUSIEURS VERTICAUX ({len(multi_verticaux)})\n")
            f.write("=" * 80 + "\n")
            f.write("(Normal si une édition contient plusieurs œuvres)\n\n")
            for entry in multi_verticaux:
                f.write(f"  {entry['id']} - {entry['nom_oeuvre']} : {entry['verticaux']} verticaux\n")
            f.write("\n")
        
        # Fiches seules - non exportées
        fiches_seules = [e for e in log_entries if not e['exporte'] and e['fiches'] > 0]
        if fiches_seules:
            f.write("=" * 80 + "\n")
            f.write("FICHES SEULES - NON EXPORTÉES\n")
            f.write("=" * 80 + "\n")
            f.write("Ces fiches n'ont pas de vertical ET/OU pas de textes associés.\n\n")
            for entry in fiches_seules:
                vertical_status = "✓" if entry['verticaux'] else "✗"
                textes_status = "✓" if entry['textes'] else "✗"
                f.write(f"  {entry['id']} - {entry['nom_oeuvre']}\n")
                f.write(f"      Vertical: {vertical_status}  |  Textes: {textes_status}\n")
            f.write("\n")
        
        # Exports sans fiche
        sans_fiche = [e for e in log_entries if e['exporte'] and 'fiche' in e['manquants']]
        if sans_fiche:
            f.write("=" * 80 + "\n")
            f.write("EXPORTS SANS FICHE (À COMPLÉTER)\n")
            f.write("=" * 80 + "\n")
            for entry in sans_fiche:
                f.write(f"  {entry['id']} - {entry['nom_oeuvre']}\n")
                f.write(f"      Chemin : {entry['export_path']}\n")
            f.write("\n")
        
        # Liste complète des exports
        exportes = [e for e in log_entries if e['exporte']]
        if exportes:
            f.write("=" * 80 + "\n")
            f.write("LISTE DES EXPORTS CRÉÉS\n")
            f.write("=" * 80 + "\n")
            
            # Libres de droits
            libres = [e for e in exportes if e['libre_de_droits']]
            if libres:
                f.write("\n[LIBRE DE DROITS]\n")
                for entry in libres:
                    status = "✓ COMPLET" if 'fiche' not in entry['manquants'] else "⚠ sans fiche"
                    f.write(f"  {entry['id']} - {entry['nom_oeuvre']} : {status}\n")
            
            # Non libres
            non_libres = [e for e in exportes if not e['libre_de_droits']]
            if non_libres:
                f.write("\n[NON LIBRE DE DROITS]\n")
                for entry in non_libres:
                    status = "✓ COMPLET" if 'fiche' not in entry['manquants'] else "⚠ sans fiche"
                    f.write(f"  {entry['id']} - {entry['nom_oeuvre']} : {status}\n")
        
        f.write("\n" + "=" * 80 + "\n")
        f.write("FIN DU RAPPORT\n")


def main(fiches_dir, verticaux_dir, textes_dir, output_dir, libres_path=None):
    """Fonction principale."""
    
    print("=" * 60)
    print("EXPORT NAKALA (VERSION PAR ÉDITION)")
    print("=" * 60)
    
    # Charger les libres de droits
    libres_set = set()
    if libres_path:
        print(f"\nChargement des libres de droits : {libres_path}")
        libres_set = load_libres_de_droits(libres_path)
    
    # Scanner les sources
    print(f"\nScan des sources...")
    print(f"  Fiches : {fiches_dir}")
    print(f"  Verticaux : {verticaux_dir}")
    print(f"  Textes : {textes_dir}")
    
    index, debug_stats = scan_sources(fiches_dir, verticaux_dir, textes_dir)
    print(f"  → {len(index)} éditions uniques trouvées")
    
    # Créer l'export
    print(f"\nCréation de l'export dans : {output_dir}")
    os.makedirs(output_dir, exist_ok=True)
    
    log_entries, stats = create_export(index, output_dir, libres_set)
    
    # Écrire le log
    log_path = os.path.join(output_dir, 'log_export.txt')
    write_log(log_entries, stats, debug_stats, log_path)
    
    # Résumé
    print("\n" + "=" * 60)
    print("RÉSUMÉ")
    print("=" * 60)
    print(f"Total trouvé : {stats['total']} éditions")
    print(f"  - Libres de droits : {stats['libres']}")
    print(f"  - Non libres de droits : {stats['non_libres']}")
    print(f"\nExports créés : {stats['exportes']}")
    print(f"  - Complets (fiche + vertical + textes) : {stats['complets']}")
    print(f"  - Sans fiche (vertical + textes) : {stats['sans_fiche']}")
    print(f"\nNon exportés (fiche seule) : {stats['fiche_seule']}")
    print(f"\nLog : {log_path}")
    
    return log_entries, stats


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Crée l'export Nakala (version par édition) en associant fiches, verticaux et textes."
    )
    parser.add_argument('--fiches', '-f', required=True, help='Dossier des fiches .docx')
    parser.add_argument('--verticaux', '-v', required=True, help='Dossier des fichiers verticaux .txt')
    parser.add_argument('--textes', '-t', required=True, help='Dossier des sous-dossiers de textes')
    parser.add_argument('--output', '-o', default='Export_Nakala', help='Dossier de sortie')
    parser.add_argument('--libres', '-l', default=None, help='Fichier listant les libres de droits')
    
    args = parser.parse_args()
    
    main(args.fiches, args.verticaux, args.textes, args.output, args.libres)
