#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Module utilitaire partagé par les scripts du dossier Nakala.

Contient les fonctions communes :
- Normalisation de noms de fichiers
- Extraction d'informations depuis les fiches .docx
- Extraction d'informations depuis les fichiers verticaux
- Chargement de la liste des libres de droits
- Parsing sécurisé des identifiants Edi-XX
- Logging centralisé
"""

import os
import re
import sys
import logging
import unicodedata
from os.path import basename


logger = logging.getLogger('nakala')


def setup_logging(verbose=False):
    """Configure le logging avec le niveau approprié.

    Args:
        verbose: Si True, active le niveau DEBUG.
    """
    level = logging.DEBUG if verbose else logging.INFO
    logging.basicConfig(
        level=level,
        format='%(levelname)s: %(message)s',
        stream=sys.stderr
    )


def parse_edi_id_number(edi_id):
    """Extrait le numéro d'un identifiant Edi-XX de manière sécurisée.

    Args:
        edi_id: Identifiant au format 'Edi-XX' (ex: 'Edi-39').

    Returns:
        Le numéro (int), ou 0 si le format est invalide.
    """
    if not edi_id or not isinstance(edi_id, str):
        return 0
    match = re.match(r'^Edi-(\d+)$', edi_id)
    if match:
        try:
            return int(match.group(1))
        except ValueError:
            return 0
    return 0


def edi_sort_key(item):
    """Clé de tri pour les identifiants Edi-XX.

    Accepte un string Edi-XX ou un objet avec attribut/clé edi_id.
    """
    if isinstance(item, str):
        return parse_edi_id_number(item)
    if hasattr(item, 'edi_id'):
        return parse_edi_id_number(item.edi_id)
    if isinstance(item, dict) and 'edi_id' in item:
        return parse_edi_id_number(item['edi_id'])
    return 0


def normalize_text(text):
    """Normalise le texte pour la comparaison : minuscules, sans accents, sans ponctuation.

    Utilise unicodedata.normalize pour une gestion complète des accents.

    Args:
        text: Texte à normaliser.

    Returns:
        Texte normalisé (minuscules, sans accents, sans ponctuation).
    """
    if not text:
        return ""
    text = text.lower()
    # Décomposition Unicode puis suppression des diacritiques
    nfkd = unicodedata.normalize('NFKD', text)
    text = ''.join(c for c in nfkd if not unicodedata.combining(c))
    # Ligatures courantes
    text = text.replace('œ', 'oe').replace('æ', 'ae')
    # Supprimer la ponctuation
    text = re.sub(r'[^\w\s]', ' ', text)
    # Normaliser les espaces
    text = ' '.join(text.split())
    return text


def normalize_filename(name, max_length=80):
    """Normalise un nom pour l'utiliser comme nom de dossier.

    Args:
        name: Nom à normaliser.
        max_length: Longueur maximale du nom.

    Returns:
        Nom normalisé, sûr pour le système de fichiers.
    """
    if not name:
        return "Inconnu"
    name = re.sub(r'[<>:"/\\|?*]', '_', name)
    name = re.sub(r'\s+', '_', name)
    name = re.sub(r'_+', '_', name)
    name = name.strip('_')
    if len(name) > max_length:
        name = name[:max_length].rstrip('_')
    return name


def extract_info_from_docx(filepath):
    """Extrait l'identifiant édition et le statut libre de droits d'une fiche .docx.

    Args:
        filepath: Chemin vers le fichier .docx.

    Returns:
        Dict avec les clés: id, titre_oeuvre, libre_de_droits.
        Retourne None en cas d'erreur.
    """
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx non installé, impossible de lire %s", filepath)
        return None

    result = {'id': None, 'titre_oeuvre': None, 'libre_de_droits': False}

    try:
        doc = Document(filepath)
        text = '\n'.join(p.text for p in doc.paragraphs)

        # Chercher l'identifiant édition
        match = re.search(r'Identifiant\s+[ée]dition\s*:\s*(Edi-\d+)', text, re.IGNORECASE)
        if match:
            result['id'] = match.group(1)

        # Chercher le titre de l'oeuvre
        titre_match = re.search(r'Titre\s*:\s*\*?([^*\n]+)\*?', text)
        if titre_match:
            result['titre_oeuvre'] = titre_match.group(1).strip()

        # Chercher si libre de droits
        if re.search(r'Libre\s+de\s+droits\s*:\s*Oui', text, re.IGNORECASE):
            result['libre_de_droits'] = True

    except Exception as e:
        logger.warning("Erreur lecture %s : %s", filepath, e)
        return None

    return result


def extract_info_from_vertical(filepath, max_lines=50):
    """Extrait l'identifiant édition d'un fichier vertical .txt.

    Args:
        filepath: Chemin vers le fichier vertical.
        max_lines: Nombre de lignes à lire pour trouver la balise <doc>.

    Returns:
        Dict avec les clés: id, source, title.
    """
    result = {'id': None, 'source': None, 'title': None}

    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            for _ in range(max_lines):
                line = f.readline()
                if not line:
                    break

                if not result['id']:
                    match = re.search(r'edition_id="(Edi-\d+)"', line)
                    if match:
                        result['id'] = match.group(1)

                if not result['source']:
                    source_match = re.search(r'source="([^"]+)"', line)
                    if source_match:
                        result['source'] = source_match.group(1)

                if not result['title']:
                    title_match = re.search(r'title="([^"]+)"', line)
                    if title_match:
                        result['title'] = title_match.group(1)

                if result['id']:
                    break

    except (IOError, OSError) as e:
        logger.warning("Erreur lecture %s : %s", filepath, e)
    except UnicodeDecodeError as e:
        logger.warning("Erreur d'encodage dans %s : %s", filepath, e)

    return result


def load_libres_de_droits(filepath):
    """Charge la liste des identifiants libres de droits depuis un fichier .docx ou .txt.

    Args:
        filepath: Chemin vers le fichier (.docx ou .txt).

    Returns:
        Set d'identifiants (ex: {'Edi-1', 'Edi-15', ...}).
    """
    libres = set()

    try:
        if filepath.endswith('.docx'):
            try:
                from docx import Document
            except ImportError:
                logger.error("python-docx requis pour lire %s", filepath)
                return libres
            doc = Document(filepath)
            text = '\n'.join(p.text for p in doc.paragraphs)
        else:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()

        matches = re.findall(r'Edi-(\d+)', text)
        for m in matches:
            libres.add(f"Edi-{m}")

        logger.info("%d identifiants libres de droits chargés", len(libres))

    except (IOError, OSError) as e:
        logger.error("Erreur chargement libres de droits depuis %s : %s", filepath, e)

    return libres


def match_textes_to_oeuvres(dirname, oeuvres_index):
    """Tente de matcher un dossier de textes à une oeuvre dans l'index.

    Args:
        dirname: Nom du dossier de textes.
        oeuvres_index: Dict {nom_fichier: oeuvre_data} où oeuvre_data a une clé 'edi_id'.

    Returns:
        La clé de l'oeuvre matchée, ou None.
    """
    # Essai 1: Match exact
    if dirname in oeuvres_index:
        return dirname

    # Essai 2: Enlever le suffixe _Edi-XX
    match_edi = re.match(r'^(.+)_(Edi-\d+)$', dirname)
    if match_edi:
        base_name = match_edi.group(1)
        edi_id = match_edi.group(2)

        if base_name in oeuvres_index:
            oeuvre = oeuvres_index[base_name]
            oeuvre_edi = oeuvre.get('edi_id') or getattr(oeuvre, 'edi_id', None)
            if oeuvre_edi == edi_id:
                return base_name

        # Match partiel du nom de base
        for key in oeuvres_index:
            if key == base_name or key.startswith(base_name) or base_name.startswith(key):
                oeuvre = oeuvres_index[key]
                oeuvre_edi = oeuvre.get('edi_id') or getattr(oeuvre, 'edi_id', None)
                if oeuvre_edi == edi_id:
                    return key

    # Essai 3: Match partiel
    for key in oeuvres_index:
        if key.startswith(dirname) or dirname.startswith(key):
            return key

    return None


def validate_path_exists(path, label="Chemin"):
    """Vérifie qu'un chemin existe et log une erreur sinon.

    Args:
        path: Chemin à vérifier.
        label: Libellé pour le message d'erreur.

    Returns:
        True si le chemin existe.
    """
    if not os.path.exists(path):
        logger.error("%s introuvable : %s", label, path)
        return False
    return True
